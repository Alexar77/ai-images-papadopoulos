# -*- coding: utf-8 -*-
"""
Concise Word document for Exercise #3 - Semantic Segmentation
"""
import os, json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r"c:\coding and uni\uni\ai-images"
NORMAL_DIR = os.path.join(BASE, "results", "ex3 with normal backround weight",
                          "experiments_20260209_234858")
LESS_DIR = os.path.join(BASE, "results", "ex3 with less backround weight",
                        "ex3", "experiments_20260213_015839")
OUTPUT = os.path.join(BASE, "Askisi_3_Simasiologiki_Tmhmatopoihsh.docx")

with open(os.path.join(NORMAL_DIR, "experiment_results.json"), encoding="utf-8") as f:
    normal_results = json.load(f)
with open(os.path.join(LESS_DIR, "experiment_results.json"), encoding="utf-8") as f:
    less_results = json.load(f)

def shade(cell, c):
    s = cell._element.get_or_add_tcPr()
    s.append(s.makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):c}))

def hd(doc, txt, lv=1):
    h = doc.add_heading(txt, level=lv)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x47,0x7A)

def im(doc, path, w=Inches(5.0), cap=None):
    if not os.path.exists(path): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=w)
    if cap:
        c = doc.add_paragraph(cap); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.runs: r.font.size = Pt(9); r.font.italic = True

def bl(doc, txt): doc.add_paragraph(txt, style='List Bullet')

def tbl(doc, cols, rows):
    t = doc.add_table(rows=1, cols=len(cols)); t.style='Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,c in enumerate(cols):
        cl = t.rows[0].cells[i]; cl.text = c; shade(cl,"1A477A")
        for p in cl.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(9)
    for rv in rows:
        row = t.add_row()
        for i,v in enumerate(rv):
            cl = row.cells[i]; cl.text = str(v)
            for p in cl.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(9)

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(11)

# ── Title ─────────────────────────────────────────────────────────────────
doc.add_paragraph("")
t=doc.add_heading("\u0386\u03c3\u03ba\u03b7\u03c3\u03b7 #3", level=0); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
st=doc.add_heading("\u03a3\u03b7\u03bc\u03b1\u03c3\u03b9\u03bf\u03bb\u03bf\u03b3\u03b9\u03ba\u03ae \u03a4\u03bc\u03b7\u03bc\u03b1\u03c4\u03bf\u03c0\u03bf\u03af\u03b7\u03c3\u03b7 (Semantic Segmentation)", level=1)
st.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph("U-Net \u03c3\u03c4\u03bf Semantic Boundaries Dataset (SBD)")
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ── 1. Introduction ──────────────────────────────────────────────────────
hd(doc, "1. \u0395\u03b9\u03c3\u03b1\u03b3\u03c9\u03b3\u03ae")
doc.add_paragraph(
    "\u0397 \u03c3\u03b7\u03bc\u03b1\u03c3\u03b9\u03bf\u03bb\u03bf\u03b3\u03b9\u03ba\u03ae \u03c4\u03bc\u03b7\u03bc\u03b1\u03c4\u03bf\u03c0\u03bf\u03af\u03b7\u03c3\u03b7 (semantic segmentation) \u03c3\u03c4\u03bf\u03c7\u03b5\u03cd\u03b5\u03b9 \u03c3\u03c4\u03b7\u03bd \u03c4\u03b1\u03be\u03b9\u03bd\u03cc\u03bc\u03b7\u03c3\u03b7 "
    "\u03ba\u03ac\u03b8\u03b5 pixel \u03bc\u03b9\u03b1\u03c2 \u03b5\u03b9\u03ba\u03cc\u03bd\u03b1\u03c2 \u03c3\u03b5 \u03ba\u03b1\u03c4\u03b7\u03b3\u03bf\u03c1\u03af\u03b1. \u03a5\u03bb\u03bf\u03c0\u03bf\u03b9\u03ae\u03b8\u03b7\u03ba\u03b5 \u03b7 \u03b1\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ae U-Net \u03ba\u03b1\u03b9 "
    "\u03b5\u03c6\u03b1\u03c1\u03bc\u03cc\u03c3\u03c4\u03b7\u03ba\u03b5 \u03c3\u03c4\u03bf SBD dataset (21 \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2: 20 \u03b1\u03bd\u03c4\u03b9\u03ba\u03b5\u03af\u03bc\u03b5\u03bd\u03b1 + background). "
    "\u0395\u03ba\u03c4\u03b5\u03bb\u03ad\u03c3\u03c4\u03b7\u03ba\u03b1\u03bd 7 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03c3\u03b5 2 \u03c3\u03b5\u03b9\u03c1\u03ad\u03c2:")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 1 (3 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1): Background weight = 1.0, 10 epochs, step scheduler")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 2 (4 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1): Background weight = 0.2, 5 epochs, cosine scheduler")

# ── 2. Methodology ───────────────────────────────────────────────────────
hd(doc, "2. \u039c\u03b5\u03b8\u03bf\u03b4\u03bf\u03bb\u03bf\u03b3\u03af\u03b1")

hd(doc, "2.1 \u0391\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ae U-Net", 2)
doc.add_paragraph(
    "\u03a4\u03bf U-Net \u03b5\u03af\u03bd\u03b1\u03b9 encoder-decoder network \u03bc\u03b5 skip connections. "
    "\u039f encoder (\u03bc\u03b5\u03af\u03c9\u03c3\u03b7 resolution \u03bc\u03ad\u03c3\u03c9 4 Down blocks: MaxPool \u2192 DoubleConv) "
    "\u03ba\u03b1\u03b9 decoder (\u03b5\u03c0\u03b1\u03bd\u03b1\u03c6\u03bf\u03c1\u03ac resolution \u03bc\u03ad\u03c3\u03c9 4 Up blocks: ConvTranspose2d \u2192 concatenation \u2192 DoubleConv). "
    "\u03a3\u03c5\u03bd\u03bf\u03bb\u03b9\u03ba\u03ad\u03c2 \u03c0\u03b1\u03c1\u03ac\u03bc\u03b5\u03c4\u03c1\u03bf\u03b9: 31.044.821.")

hd(doc, "2.2 \u03a1\u03cc\u03bb\u03bf\u03c2 Background Weight", 2)
doc.add_paragraph(
    "\u03a4\u03bf background \u03ba\u03b1\u03c4\u03b1\u03bb\u03b1\u03bc\u03b2\u03ac\u03bd\u03b5\u03b9 \u03c4\u03b7 \u03c3\u03c5\u03bd\u03c4\u03c1\u03b9\u03c0\u03c4\u03b9\u03ba\u03ae \u03c0\u03bb\u03b5\u03b9\u03bf\u03bd\u03cc\u03c4\u03b7\u03c4\u03b1 \u03c4\u03c9\u03bd pixels. "
    "\u039c\u03b5 weight=1.0 \u03c4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03bc\u03c0\u03bf\u03c1\u03b5\u03af \u03bd\u03b1 \u03c0\u03b5\u03c4\u03cd\u03c7\u03b5\u03b9 \u03c5\u03c8\u03b7\u03bb\u03ae pixel accuracy \u03c0\u03c1\u03bf\u03b2\u03bb\u03ad\u03c0\u03bf\u03bd\u03c4\u03b1\u03c2 "
    "background \u03c0\u03b1\u03bd\u03c4\u03bf\u03cd. \u039c\u03b5\u03b9\u03ce\u03bd\u03bf\u03bd\u03c4\u03b1\u03c2 \u03c4\u03bf \u03c3\u03c4\u03bf 0.2, \u03b7 loss \u03c4\u03b9\u03bc\u03c9\u03c1\u03b5\u03af \u03bb\u03b9\u03b3\u03cc\u03c4\u03b5\u03c1\u03bf \u03c4\u03b1 \u03bb\u03ac\u03b8\u03b7 \u03c3\u03c4\u03bf "
    "background, \u03b1\u03bd\u03b1\u03b3\u03ba\u03ac\u03b6\u03bf\u03bd\u03c4\u03b1\u03c2 \u03c4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03bd\u03b1 \u03b5\u03c3\u03c4\u03b9\u03ac\u03c3\u03b5\u03b9 \u03c3\u03c4\u03b9\u03c2 foreground \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2 "
    "\u2014 \u03b1\u03bb\u03bb\u03ac \u03bc\u03c0\u03bf\u03c1\u03b5\u03af \u03bd\u03b1 \u03b1\u03c0\u03bf\u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03bf\u03c0\u03bf\u03b9\u03ae\u03c3\u03b5\u03b9 \u03c4\u03b7\u03bd \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7.")

# ── 3. Results ────────────────────────────────────────────────────────────
hd(doc, "3. \u0391\u03c0\u03bf\u03c4\u03b5\u03bb\u03ad\u03c3\u03bc\u03b1\u03c4\u03b1 \u03a0\u03b5\u03b9\u03c1\u03b1\u03bc\u03ac\u03c4\u03c9\u03bd")

hd(doc, "3.1 \u03a3\u03c5\u03bd\u03bf\u03c0\u03c4\u03b9\u03ba\u03cc\u03c2 \u03a0\u03af\u03bd\u03b1\u03ba\u03b1\u03c2", 2)

cols = ["\u03a0\u03b5\u03af\u03c1\u03b1\u03bc\u03b1", "BG Wt", "Optimizer", "LR", "Epochs", "Scheduler",
        "mIoU(%)", "PixAcc(%)", "\u03a7\u03c1\u03cc\u03bd\u03bf\u03c2(min)"]
rows = []
for exp in normal_results.values():
    hp = exp.get("hyperparameters", exp.get("config", {}))
    ep = len(exp["history"]["train_loss"])
    rows.append([exp["name"], "1.0", hp.get("optimizer","").upper(),
                 hp.get("learning_rate",""), ep, hp.get("scheduler",""),
                 f'{exp["val_miou"]:.2f}', f'{exp["val_pixel_acc"]:.2f}',
                 f'{exp["total_time"]/60:.1f}'])
for exp in less_results.values():
    hp = exp.get("hyperparameters", exp.get("config", {}))
    ep = len(exp["history"]["train_loss"])
    rows.append([exp["name"], hp.get("background_weight","0.2"),
                 hp.get("optimizer","").upper(), hp.get("learning_rate",""),
                 ep, hp.get("scheduler",""),
                 f'{exp["val_miou"]:.2f}', f'{exp["val_pixel_acc"]:.2f}',
                 f'{exp["total_time"]/60:.1f}'])
rows.sort(key=lambda r: float(r[6]), reverse=True)
tbl(doc, cols, rows)

doc.add_paragraph("")

# ── 3.2 Normal BG ────────────────────────────────────────────────────────
hd(doc, "3.2 \u03a3\u03b5\u03b9\u03c1\u03ac 1: \u039a\u03b1\u03bd\u03bf\u03bd\u03b9\u03ba\u03cc Background Weight (=1.0)", 2)
doc.add_paragraph(
    "3 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03bc\u03b5 10 epochs, step scheduler, \u03cc\u03bb\u03b5\u03c2 \u03bf\u03b9 \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2 \u03b9\u03c3\u03cc\u03c4\u03b9\u03bc\u03b1 \u03c3\u03c4\u03b7 loss function:")

im(doc, os.path.join(NORMAL_DIR, "report_01_miou_overview.png"),
   Inches(4.5), "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 1: \u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 mIoU \u2014 \u039a\u03b1\u03bd\u03bf\u03bd\u03b9\u03ba\u03cc BG Weight")

im(doc, os.path.join(NORMAL_DIR, "report_07_best_segmentation_panel.png"),
   Inches(5.2), "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 2: \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03b1\u03c0\u03bf\u03c4\u03ad\u03bb\u03b5\u03c3\u03bc\u03b1 \u03c4\u03bc\u03b7\u03bc\u03b1\u03c4\u03bf\u03c0\u03bf\u03af\u03b7\u03c3\u03b7\u03c2 (BG=1.0)")

doc.add_paragraph("\u03a3\u03c7\u03bf\u03bb\u03b9\u03b1\u03c3\u03bc\u03cc\u03c2:")
bl(doc, "\u039f AdamW (lr=0.0003) \u03c0\u03ad\u03c4\u03c5\u03c7\u03b5 \u03c4\u03bf \u03ba\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf mIoU (13.91%) \u03c7\u03ac\u03c1\u03b7 \u03c3\u03c4\u03b7\u03bd adaptive \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03ae \u03c4\u03bf\u03c5.")
bl(doc, "\u039f SGD (lr=0.0003) \u03c3\u03c5\u03bd\u03ad\u03ba\u03bb\u03b9\u03bd\u03b5 \u03c0\u03b9\u03bf \u03b1\u03c1\u03b3\u03ac \u03ba\u03b1\u03b9 \u03ad\u03ba\u03b1\u03bd\u03b5 early stop \u03c3\u03c4\u03b1 7 epochs (mIoU=10.25%).")
bl(doc, "\u039c\u03b5\u03b3\u03ac\u03bb\u03bf LR (0.001) \u03ad\u03b4\u03c9\u03c3\u03b5 \u03b5\u03bb\u03b1\u03c6\u03c1\u03ce\u03c2 \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf mIoU (12.63%) \u03b1\u03c0\u03cc \u03c4\u03bf 0.0003.")
bl(doc, "\u0397 \u03c5\u03c8\u03b7\u03bb\u03ae Pixel Accuracy (71-75%) \u03bf\u03c6\u03b5\u03af\u03bb\u03b5\u03c4\u03b1\u03b9 \u03c3\u03c4\u03b7\u03bd \u03ba\u03c5\u03c1\u03b9\u03b1\u03c1\u03c7\u03af\u03b1 \u03c4\u03bf\u03c5 background \u03c3\u03c4\u03b1 pixels.")

# ── 3.3 Less BG ──────────────────────────────────────────────────────────
hd(doc, "3.3 \u03a3\u03b5\u03b9\u03c1\u03ac 2: \u039c\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03bf Background Weight (=0.2)", 2)
doc.add_paragraph(
    "4 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03bc\u03b5 5 epochs, cosine scheduler, "
    "\u03b2\u03ac\u03c1\u03bf\u03c2 background=0.2 \u03ce\u03c3\u03c4\u03b5 \u03c4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03bd\u03b1 \u03b5\u03c3\u03c4\u03b9\u03ac\u03c3\u03b5\u03b9 \u03c3\u03c4\u03b9\u03c2 foreground \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2:")

im(doc, os.path.join(LESS_DIR, "report_01_miou_overview.png"),
   Inches(4.5), "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 3: \u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 mIoU \u2014 \u039c\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03bf BG Weight (0.2)")

im(doc, os.path.join(LESS_DIR, "report_07_best_segmentation_panel.png"),
   Inches(5.2), "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 4: \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03b1\u03c0\u03bf\u03c4\u03ad\u03bb\u03b5\u03c3\u03bc\u03b1 \u03c4\u03bc\u03b7\u03bc\u03b1\u03c4\u03bf\u03c0\u03bf\u03af\u03b7\u03c3\u03b7\u03c2 (BG=0.2)")

doc.add_paragraph("\u03a3\u03c7\u03bf\u03bb\u03b9\u03b1\u03c3\u03bc\u03cc\u03c2:")
bl(doc, "\u038c\u03bb\u03b1 \u03c4\u03b1 mIoU \u03b5\u03af\u03bd\u03b1\u03b9 \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03b1 (3-7.5% vs 10-14%) \u03bb\u03cc\u03b3\u03c9 \u03bb\u03b9\u03b3\u03cc\u03c4\u03b5\u03c1\u03c9\u03bd epochs (5 vs 10) "
   "\u03ba\u03b1\u03b9 \u03b1\u03c0\u03bf\u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03bf\u03c0\u03bf\u03af\u03b7\u03c3\u03b7\u03c2 \u03b1\u03c0\u03cc \u03c4\u03b1 \u03b1\u03bd\u03b9\u03c3\u03cc\u03c1\u03c1\u03bf\u03c0\u03b1 \u03b2\u03ac\u03c1\u03b7.")
bl(doc, "\u039f AdamW \u03c0\u03b1\u03c1\u03b1\u03bc\u03ad\u03bd\u03b5\u03b9 \u03c5\u03c0\u03ad\u03c1\u03c4\u03b5\u03c1\u03bf\u03c2 (6.75% vs 3.02% \u03c3\u03b5 \u03c3\u03c7\u03ad\u03c3\u03b7 \u03bc\u03b5 SGD).")
bl(doc, "\u03a4\u03bf \u03bc\u03b9\u03ba\u03c1\u03cc\u03c4\u03b5\u03c1\u03bf LR (0.0001) \u03b4\u03af\u03bd\u03b5\u03b9 \u03c4\u03bf \u03ba\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03b1\u03c0\u03bf\u03c4\u03ad\u03bb\u03b5\u03c3\u03bc\u03b1 (7.56%) \u2014 "
   "\u03b7 \u03c0\u03b9\u03bf \u03b1\u03c1\u03b3\u03ae \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03b7 \u03b2\u03bf\u03b7\u03b8\u03ac \u03cc\u03c4\u03b1\u03bd \u03c4\u03b1 \u03b2\u03ac\u03c1\u03b7 \u03b5\u03af\u03bd\u03b1\u03b9 \u03b1\u03bd\u03b9\u03c3\u03cc\u03c1\u03c1\u03bf\u03c0\u03b1.")
bl(doc, "\u0397 Pixel Accuracy \u03ad\u03c0\u03b5\u03c3\u03b5 \u03c3\u03c4\u03bf 67-71%, \u03b3\u03b9\u03b1\u03c4\u03af \u03c4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03b4\u03b5\u03bd \u03c4\u03b9\u03bc\u03c9\u03c1\u03b5\u03af\u03c4\u03b1\u03b9 "
   "\u03b1\u03c1\u03ba\u03b5\u03c4\u03ac \u03b3\u03b9\u03b1 \u03bb\u03ac\u03b8\u03b7 \u03c3\u03c4\u03bf background.")

# ── 4. Comparative Analysis ──────────────────────────────────────────────
hd(doc, "4. \u03a3\u03c5\u03b3\u03ba\u03c1\u03b9\u03c4\u03b9\u03ba\u03ae \u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7")

hd(doc, "4.1 \u039a\u03b1\u03bd\u03bf\u03bd\u03b9\u03ba\u03cc vs \u039c\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03bf Background Weight", 2)

doc.add_paragraph(
    "\u039a\u03b1\u03bd\u03bf\u03bd\u03b9\u03ba\u03cc BG Weight (=1.0):")
bl(doc, "\u03a3\u03c4\u03b1\u03b8\u03b5\u03c1\u03ae \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7, \u03c5\u03c8\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf mIoU (13.91%) \u03ba\u03b1\u03b9 Pixel Accuracy (74.79%).")
bl(doc, "\u03a4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03b5\u03c3\u03c4\u03b9\u03ac\u03b6\u03b5\u03b9 \u03c0\u03b5\u03c1\u03b9\u03c3\u03c3\u03cc\u03c4\u03b5\u03c1\u03bf \u03c3\u03c4\u03bf background \u03bb\u03cc\u03b3\u03c9 \u03c4\u03b7\u03c2 \u03ba\u03c5\u03c1\u03b9\u03b1\u03c1\u03c7\u03af\u03b1\u03c2 \u03c4\u03bf\u03c5.")

doc.add_paragraph(
    "\u039c\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03bf BG Weight (=0.2):")
bl(doc, "\u0397 \u03c4\u03b9\u03bc\u03c9\u03c1\u03af\u03b1 \u03b3\u03b9\u03b1 \u03bb\u03ac\u03b8\u03b7 \u03c3\u03c4\u03bf background \u03c0\u03ad\u03c6\u03c4\u03b5\u03b9 \u03c3\u03c4\u03bf 20%, "
   "\u03b1\u03bd\u03b1\u03b3\u03ba\u03ac\u03b6\u03bf\u03bd\u03c4\u03b1\u03c2 \u03c4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03bd\u03b1 \u03bc\u03ac\u03b8\u03b5\u03b9 \u03c4\u03b9\u03c2 foreground \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2.")
bl(doc, "\u03a7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03b1 \u03b1\u03c0\u03bf\u03c4\u03b5\u03bb\u03ad\u03c3\u03bc\u03b1\u03c4\u03b1 (7.56%) \u03bb\u03cc\u03b3\u03c9 \u03bb\u03b9\u03b3\u03cc\u03c4\u03b5\u03c1\u03c9\u03bd epochs \u03ba\u03b1\u03b9 \u03b1\u03c3\u03c4\u03ac\u03b8\u03b5\u03b9\u03b1\u03c2.")
bl(doc, "\u039c\u03b5 \u03c0\u03b5\u03c1\u03b9\u03c3\u03c3\u03cc\u03c4\u03b5\u03c1\u03b5\u03c2 epochs \u03b1\u03bd\u03b1\u03bc\u03ad\u03bd\u03b5\u03c4\u03b1\u03b9 \u03bd\u03b1 \u03c5\u03c0\u03b5\u03c1\u03c4\u03b5\u03c1\u03ae\u03c3\u03b5\u03b9 \u03c3\u03c4\u03bf foreground mIoU.")

hd(doc, "4.2 \u0395\u03c0\u03af\u03b4\u03c1\u03b1\u03c3\u03b7 Optimizer", 2)
bl(doc, "\u039f AdamW \u03c5\u03c0\u03b5\u03c1\u03c4\u03b5\u03c1\u03b5\u03af \u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03ac \u03ad\u03bd\u03b1\u03bd\u03c4\u03b9 \u03c4\u03bf\u03c5 SGD \u03c3\u03b5 \u03cc\u03bb\u03b1 \u03c4\u03b1 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1.")
bl(doc, "\u0397 \u03b4\u03b9\u03b1\u03c6\u03bf\u03c1\u03ac \u03b5\u03af\u03bd\u03b1\u03b9 \u03c0\u03b9\u03bf \u03ad\u03bd\u03c4\u03bf\u03bd\u03b7 \u03bc\u03b5 \u03bb\u03b9\u03b3\u03cc\u03c4\u03b5\u03c1\u03b1 epochs: AdamW 6.75% vs SGD 3.02% (BG=0.2).")
bl(doc, "\u039f SGD \u03c7\u03c1\u03b5\u03b9\u03ac\u03b6\u03b5\u03c4\u03b1\u03b9 \u03c0\u03b5\u03c1\u03b9\u03c3\u03c3\u03cc\u03c4\u03b5\u03c1\u03b1 epochs \u03ba\u03b1\u03b9 \u03c0\u03c1\u03bf\u03c3\u03b5\u03ba\u03c4\u03b9\u03ba\u03cc\u03c4\u03b5\u03c1\u03bf LR tuning.")

hd(doc, "4.3 \u0395\u03c0\u03af\u03b4\u03c1\u03b1\u03c3\u03b7 Learning Rate", 2)
bl(doc, "BG=1.0: lr=0.0003 > lr=0.001 (13.91% vs 12.63%).")
bl(doc, "BG=0.2: lr=0.0001 > lr=0.0003 > lr=0.001 (7.56% > 6.75% > 5.80%).")
bl(doc, "\u038c\u03c3\u03bf \u03c0\u03b9\u03bf \u03b1\u03c3\u03c4\u03b1\u03b8\u03ae\u03c2 \u03b7 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7, \u03c4\u03cc\u03c3\u03bf \u03c0\u03b9\u03bf \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc LR \u03c7\u03c1\u03b5\u03b9\u03ac\u03b6\u03b5\u03c4\u03b1\u03b9.")

# ── 5. Conclusions ────────────────────────────────────────────────────────
hd(doc, "5. \u03a3\u03c5\u03bc\u03c0\u03b5\u03c1\u03ac\u03c3\u03bc\u03b1\u03c4\u03b1")
bl(doc, "\u03a4\u03bf U-Net \u03b5\u03af\u03bd\u03b1\u03b9 \u03b1\u03c0\u03bf\u03c4\u03b5\u03bb\u03b5\u03c3\u03bc\u03b1\u03c4\u03b9\u03ba\u03ae \u03b1\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ae \u03b3\u03b9\u03b1 semantic segmentation \u03bc\u03b5 skip connections.")
bl(doc, "\u039f AdamW \u03b5\u03af\u03bd\u03b1\u03b9 \u03bf \u03ba\u03b1\u03c4\u03b1\u03bb\u03bb\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf\u03c2 optimizer, \u03b9\u03b4\u03b9\u03b1\u03af\u03c4\u03b5\u03c1\u03b1 \u03bc\u03b5 \u03c0\u03b5\u03c1\u03b9\u03bf\u03c1\u03b9\u03c3\u03bc\u03ad\u03bd\u03bf \u03b1\u03c1\u03b9\u03b8\u03bc\u03cc epochs.")
bl(doc, "\u0397 \u03bc\u03b5\u03af\u03c9\u03c3\u03b7 \u03c4\u03bf\u03c5 background weight \u03c3\u03c4\u03bf\u03c7\u03b5\u03cd\u03b5\u03b9 \u03c3\u03c4\u03b9\u03c2 foreground \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2 \u03b1\u03bb\u03bb\u03ac \u03b1\u03c0\u03b1\u03b9\u03c4\u03b5\u03af \u03c0\u03b5\u03c1\u03b9\u03c3\u03c3\u03cc\u03c4\u03b5\u03c1\u03b5\u03c2 epochs.")
bl(doc, "\u03a7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf LR (0.0001-0.0003) \u03bf\u03b4\u03b7\u03b3\u03b5\u03af \u03c3\u03b5 \u03ba\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03b1 \u03b1\u03c0\u03bf\u03c4\u03b5\u03bb\u03ad\u03c3\u03bc\u03b1\u03c4\u03b1 \u03c3\u03b5 \u03b1\u03c3\u03c4\u03b1\u03b8\u03b5\u03af\u03c2 \u03c3\u03c5\u03bd\u03b8\u03ae\u03ba\u03b5\u03c2.")
bl(doc, "\u03a4\u03bf Mean IoU (\u03c3\u03c4\u03b9\u03c2 foreground \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2) \u03b5\u03af\u03bd\u03b1\u03b9 \u03c0\u03b9\u03bf \u03b1\u03be\u03b9\u03cc\u03c0\u03b9\u03c3\u03c4\u03b7 \u03bc\u03b5\u03c4\u03c1\u03b9\u03ba\u03ae \u03b1\u03c0\u03cc \u03c4\u03b7\u03bd Pixel Accuracy.")

doc.save(OUTPUT)
print(f"\n[OK] Saved: {OUTPUT}")
print(f"  Size: {os.path.getsize(OUTPUT)/(1024*1024):.1f} MB")
