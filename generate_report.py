# -*- coding: utf-8 -*-
"""Combined Word report for Exercises 1, 2, 3"""
import os, json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r"c:\coding and uni\uni\ai-images"
EX1 = os.path.join(BASE, "results", "ex1", "experiments_20260208_192243")
EX2 = os.path.join(BASE, "results", "ex2", "experiments_20260208_201208")
EX3N = os.path.join(BASE, "results", "ex3 with normal backround weight", "experiments_20260209_234858")
EX3L = os.path.join(BASE, "results", "ex3 with less backround weight", "ex3", "experiments_20260213_015839")
EX4 = os.path.join(BASE, "results", "ex4")
EX5 = os.path.join(BASE, "results", "ex5")
OUTPUT = os.path.join(BASE, "Report_Exercises_1_2_3_4_5.docx")

def load_json(d):
    with open(os.path.join(d, "experiment_results.json"), encoding="utf-8") as f:
        return json.load(f)

ex1 = load_json(EX1); ex2 = load_json(EX2); ex3n = load_json(EX3N); ex3l = load_json(EX3L)

# Helpers
def shade(cell, c):
    s = cell._element.get_or_add_tcPr()
    s.append(s.makeelement(qn('w:shd'), {qn('w:val'):'clear',qn('w:color'):'auto',qn('w:fill'):c}))

def hd(doc, txt, lv=1):
    h = doc.add_heading(txt, level=lv)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x47,0x7A)

def im(doc, path, w=Inches(4.8), cap=None):
    if not os.path.exists(path): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=w)
    if cap:
        c = doc.add_paragraph(cap); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.runs: r.font.size = Pt(9); r.font.italic = True

def bl(doc, txt): doc.add_paragraph(txt, style='List Bullet')

def tbl(doc, cols, rows):
    t = doc.add_table(rows=1, cols=len(cols)); t.style='Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,c in enumerate(cols):
        cl = t.rows[0].cells[i]; cl.text = c; shade(cl,"1A477A")
        for p in cl.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(8)
    for rv in rows:
        row = t.add_row()
        for i,v in enumerate(rv):
            cl = row.cells[i]; cl.text = str(v)
            for p in cl.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.size = Pt(8)

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(11)

# ══════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph("")
doc.add_paragraph("")
t=doc.add_heading("\u0391\u03c3\u03ba\u03ae\u03c3\u03b5\u03b9\u03c2 #1, #2, #3, #4, #5", level=0)
t.alignment=WD_ALIGN_PARAGRAPH.CENTER
st=doc.add_heading("\u0395\u03c0\u03b5\u03be\u03b5\u03c1\u03b3\u03b1\u03c3\u03af\u03b1 \u0395\u03b9\u03ba\u03cc\u03bd\u03b1\u03c2 \u03bc\u03b5 \u0392\u03b1\u03b8\u03b9\u03ac \u039c\u03ac\u03b8\u03b7\u03c3\u03b7", level=1)
st.alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph("CNN, Transfer Learning, Semantic Segmentation, Object Detection, CNN vs ViT")
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# EXERCISE 1
# ══════════════════════════════════════════════════════════════════════════
hd(doc, "\u0386\u03c3\u03ba\u03b7\u03c3\u03b7 #1: \u0392\u03b5\u03bb\u03c4\u03b9\u03c3\u03c4\u03bf\u03c0\u03bf\u03af\u03b7\u03c3\u03b7 \u0391\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ae\u03c2 \u03ba\u03b1\u03b9 \u03a5\u03c0\u03b5\u03c1-\u03c0\u03b1\u03c1\u03b1\u03bc\u03ad\u03c4\u03c1\u03c9\u03bd")

hd(doc, "1.1 \u0395\u03b9\u03c3\u03b1\u03b3\u03c9\u03b3\u03ae", 2)
doc.add_paragraph(
    "\u03a5\u03bb\u03bf\u03c0\u03bf\u03b9\u03ae\u03b8\u03b7\u03ba\u03b5 \u03ad\u03bd\u03b1 Vanilla CNN \u03bc\u03b5 4 \u03c3\u03c5\u03bd\u03b5\u03bb\u03b9\u03ba\u03c4\u03b9\u03ba\u03ac \u03c3\u03c4\u03c1\u03ce\u03bc\u03b1\u03c4\u03b1 (964.516 \u03c0\u03b1\u03c1\u03ac\u03bc\u03b5\u03c4\u03c1\u03bf\u03b9) "
    "\u03b3\u03b9\u03b1 \u03c4\u03b1\u03be\u03b9\u03bd\u03cc\u03bc\u03b7\u03c3\u03b7 \u03c3\u03c4\u03bf CIFAR-100 (100 \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2, 32x32 pixels). "
    "\u03a4\u03bf backbone \u03b1\u03c0\u03bf\u03c4\u03b5\u03bb\u03b5\u03af\u03c4\u03b1\u03b9 \u03b1\u03c0\u03cc Conv2d \u2192 BatchNorm \u2192 ReLU \u2192 MaxPool blocks, "
    "\u03b1\u03ba\u03bf\u03bb\u03bf\u03c5\u03b8\u03bf\u03cd\u03bc\u03b5\u03bd\u03bf \u03b1\u03c0\u03cc fully connected classifier. "
    "\u0395\u03ba\u03c4\u03b5\u03bb\u03ad\u03c3\u03c4\u03b7\u03ba\u03b1\u03bd 8 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03c3\u03b5 3 \u03c3\u03b5\u03b9\u03c1\u03ad\u03c2 (30 epochs, cosine scheduler, batch size 128):")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 1 \u2014 \u03a3\u03c5\u03bd\u03ac\u03c1\u03c4\u03b7\u03c3\u03b7 \u03ba\u03cc\u03c3\u03c4\u03bf\u03c5\u03c2: CrossEntropy (baseline: SGD, lr=0.01)")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 2 \u2014 Optimizers: SGD, Adam, AdamW, RMSprop (\u03cc\u03bb\u03bf\u03b9 \u03bc\u03b5 lr=0.01)")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 3 \u2014 Learning rates: 0.001, 0.01, 0.1 (\u03cc\u03bb\u03b1 \u03bc\u03b5 SGD)")

hd(doc, "1.2 \u03a3\u03c5\u03bd\u03bf\u03c0\u03c4\u03b9\u03ba\u03cc\u03c2 \u03a0\u03af\u03bd\u03b1\u03ba\u03b1\u03c2 \u0391\u03c0\u03bf\u03c4\u03b5\u03bb\u03b5\u03c3\u03bc\u03ac\u03c4\u03c9\u03bd", 2)
cols1 = ["\u03a0\u03b5\u03af\u03c1\u03b1\u03bc\u03b1", "Optimizer", "LR", "Epochs", "Train(%)", "Test(%)", "\u03a7\u03c1\u03cc\u03bd\u03bf\u03c2(min)"]
rows1 = []
for exp in ex1.values():
    hp = exp.get("hyperparameters", exp.get("config", {}))
    ep = len(exp["history"]["train_loss"])
    rows1.append([exp["name"], hp.get("optimizer","").upper(), hp.get("learning_rate",""),
                  ep, f'{exp["train_acc"]:.1f}', f'{exp["test_acc"]:.2f}',
                  f'{exp["total_time"]/60:.1f}'])
rows1.sort(key=lambda r: float(r[5]), reverse=True)
tbl(doc, cols1, rows1)
doc.add_paragraph("")

im(doc, os.path.join(EX1, "report_01_accuracy_overview.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 1.1: \u0395\u03c0\u03b9\u03c3\u03ba\u03cc\u03c0\u03b7\u03c3\u03b7 \u03b1\u03ba\u03c1\u03af\u03b2\u03b5\u03b9\u03b1\u03c2 \u03cc\u03bb\u03c9\u03bd \u03c4\u03c9\u03bd \u03c0\u03b5\u03b9\u03c1\u03b1\u03bc\u03ac\u03c4\u03c9\u03bd")

# ── 1.3 Optimizer comparison ─────────────────────────────────────────────
hd(doc, "1.3 \u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 Optimizers", 2)
doc.add_paragraph(
    "\u038c\u03bb\u03bf\u03b9 \u03bf\u03b9 optimizers \u03b4\u03bf\u03ba\u03b9\u03bc\u03ac\u03c3\u03c4\u03b7\u03ba\u03b1\u03bd \u03bc\u03b5 lr=0.01. \u03a4\u03b1 \u03b1\u03c0\u03bf\u03c4\u03b5\u03bb\u03ad\u03c3\u03bc\u03b1\u03c4\u03b1 \u03b4\u03b5\u03af\u03c7\u03bd\u03bf\u03c5\u03bd \u03cc\u03c4\u03b9 "
    "\u03b7 \u03b5\u03c0\u03b9\u03bb\u03bf\u03b3\u03ae optimizer \u03ad\u03c7\u03b5\u03b9 \u03c4\u03b5\u03c1\u03ac\u03c3\u03c4\u03b9\u03b1 \u03b5\u03c0\u03af\u03b4\u03c1\u03b1\u03c3\u03b7 \u03c3\u03c4\u03b7\u03bd \u03b1\u03c0\u03cc\u03b4\u03bf\u03c3\u03b7:")

im(doc, os.path.join(EX1, "report_02_optimizer_curves.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 1.2: \u039a\u03b1\u03bc\u03c0\u03cd\u03bb\u03b5\u03c2 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7\u03c2 \u03b1\u03bd\u03ac optimizer")

doc.add_paragraph("\u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7:")
bl(doc, "SGD (50.48%): \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf\u03c2 optimizer \u03bc\u03b5 lr=0.01. \u039f SGD \u03c7\u03c1\u03b7\u03c3\u03b9\u03bc\u03bf\u03c0\u03bf\u03b9\u03b5\u03af \u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03cc gradient scaling, "
   "\u03c0\u03bf\u03c5 \u03c3\u03b5 \u03c3\u03c5\u03bd\u03b4\u03c5\u03b1\u03c3\u03bc\u03cc \u03bc\u03b5 \u03c4\u03bf cosine scheduler \u03b5\u03c0\u03b9\u03c4\u03c1\u03ad\u03c0\u03b5\u03b9 \u03bf\u03bc\u03b1\u03bb\u03ae \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03b7. "
   "\u0397 \u03ba\u03b1\u03bc\u03c0\u03cd\u03bb\u03b7 \u03b1\u03ba\u03c1\u03af\u03b2\u03b5\u03b9\u03b1\u03c2 \u03b1\u03bd\u03b5\u03b2\u03b1\u03af\u03bd\u03b5\u03b9 \u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03ac \u03c3\u03b5 \u03cc\u03bb\u03b1 \u03c4\u03b1 30 epochs.")
bl(doc, "AdamW (12.72%): \u0391\u03c0\u03bf\u03c4\u03c5\u03c7\u03af\u03b1 \u03bb\u03cc\u03b3\u03c9 \u03c5\u03c8\u03b7\u03bb\u03bf\u03cd lr. \u039f\u03b9 adaptive optimizers (\u03cc\u03c0\u03c9\u03c2 Adam/AdamW) "
   "\u03c0\u03c1\u03bf\u03c3\u03b1\u03c1\u03bc\u03cc\u03b6\u03bf\u03c5\u03bd \u03c4\u03bf LR \u03b1\u03bd\u03ac \u03c0\u03b1\u03c1\u03ac\u03bc\u03b5\u03c4\u03c1\u03bf, \u03ac\u03c1\u03b1 \u03bc\u03b5 lr=0.01 \u03c4\u03bf \u03c0\u03c1\u03b1\u03b3\u03bc\u03b1\u03c4\u03b9\u03ba\u03cc LR \u03b3\u03af\u03bd\u03b5\u03c4\u03b1\u03b9 "
   "\u03c0\u03bf\u03bb\u03cd \u03c5\u03c8\u03b7\u03bb\u03cc \u03b3\u03b9\u03b1 \u03ad\u03bd\u03b1 \u03bc\u03b9\u03ba\u03c1\u03cc CNN, \u03c0\u03c1\u03bf\u03ba\u03b1\u03bb\u03ce\u03bd\u03c4\u03b1\u03c2 \u03b1\u03c3\u03c4\u03ac\u03b8\u03b5\u03b9\u03b1 \u03ba\u03b1\u03b9 \u03b1\u03c1\u03b3\u03ae \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03b7.")
bl(doc, "Adam (1.0%): \u03a0\u03bb\u03ae\u03c1\u03b7\u03c2 \u03b1\u03c0\u03bf\u03c4\u03c5\u03c7\u03af\u03b1 \u2014 \u03c4\u03bf loss \u03c0\u03b1\u03c1\u03ad\u03bc\u03b5\u03b9\u03bd\u03b5 \u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03cc (~4.605) \u03b3\u03b9\u03b1 30 epochs. "
   "\u039c\u03b5 lr=0.01 \u03bf Adam \u03c5\u03c0\u03b5\u03c1\u03c0\u03ae\u03b4\u03b7\u03c3\u03b5 \u03c4\u03bf \u03b5\u03bb\u03ac\u03c7\u03b9\u03c3\u03c4\u03bf \u03ba\u03b1\u03b9 \u03c0\u03b1\u03b3\u03b9\u03b4\u03b5\u03cd\u03c4\u03b7\u03ba\u03b5 \u03c3\u03b5 \u03c4\u03c5\u03c7\u03b1\u03af\u03b1 \u03c0\u03c1\u03cc\u03b2\u03bb\u03b5\u03c8\u03b7 "
   "(1% = \u03c4\u03c5\u03c7\u03b1\u03af\u03bf \u03b3\u03b9\u03b1 100 \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2). \u0391\u03c5\u03c4\u03cc \u03b4\u03b5\u03af\u03c7\u03bd\u03b5\u03b9 \u03cc\u03c4\u03b9 \u03bf Adam \u03b1\u03c0\u03b1\u03b9\u03c4\u03b5\u03af lr \u2264 0.001 \u03b3\u03b9\u03b1 \u03bc\u03b9\u03ba\u03c1\u03ac CNNs.")
bl(doc, "RMSprop (6.51%): \u0391\u03c3\u03c4\u03b1\u03b8\u03ae\u03c2 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7 \u03bc\u03b5 \u03ad\u03bd\u03c4\u03bf\u03bd\u03b5\u03c2 \u03b4\u03b9\u03b1\u03ba\u03c5\u03bc\u03ac\u03bd\u03c3\u03b5\u03b9\u03c2 \u03c3\u03c4\u03bf loss "
   "(\u03c4\u03bf \u03c0\u03c1\u03ce\u03c4\u03bf epoch \u03b5\u03af\u03c7\u03b5 loss=5355!). \u0388\u03ba\u03b1\u03bd\u03b5 early stop \u03c3\u03c4\u03b1 13 epochs. "
   "\u039f RMSprop \u03b4\u03b5\u03bd \u03c7\u03c1\u03b7\u03c3\u03b9\u03bc\u03bf\u03c0\u03bf\u03b9\u03b5\u03af momentum \u03ba\u03b1\u03b9 \u03b5\u03af\u03bd\u03b1\u03b9 \u03b5\u03c5\u03b1\u03af\u03c3\u03b8\u03b7\u03c4\u03bf\u03c2 \u03c3\u03b5 \u03c5\u03c8\u03b7\u03bb\u03cc lr.")

im(doc, os.path.join(EX1, "report_04_generalization_gap.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 1.3: Generalization gap (train vs test accuracy)")

# ── 1.4 Learning Rate comparison ─────────────────────────────────────────
hd(doc, "1.4 \u0395\u03c0\u03af\u03b4\u03c1\u03b1\u03c3\u03b7 Learning Rate", 2)
doc.add_paragraph(
    "\u0397 \u03c3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 \u03c4\u03c1\u03b9\u03ce\u03bd learning rates \u03bc\u03b5 SGD \u03b4\u03b5\u03af\u03c7\u03bd\u03b5\u03b9 \u03cc\u03c4\u03b9 \u03c4\u03bf LR "
    "\u03b5\u03af\u03bd\u03b1\u03b9 \u03b7 \u03c0\u03b9\u03bf \u03ba\u03c1\u03af\u03c3\u03b9\u03bc\u03b7 \u03c5\u03c0\u03b5\u03c1-\u03c0\u03b1\u03c1\u03ac\u03bc\u03b5\u03c4\u03c1\u03bf\u03c2:")

im(doc, os.path.join(EX1, "report_03_lr_train_loss_curves.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 1.4: \u039a\u03b1\u03bc\u03c0\u03cd\u03bb\u03b5\u03c2 training loss \u03b1\u03bd\u03ac learning rate")

doc.add_paragraph("\u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7:")
bl(doc, "lr=0.1 (56.43%): \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03b1\u03c0\u03bf\u03c4\u03ad\u03bb\u03b5\u03c3\u03bc\u03b1. \u03a4\u03bf \u03c5\u03c8\u03b7\u03bb\u03cc lr \u03b5\u03c0\u03b9\u03c4\u03c1\u03ad\u03c0\u03b5\u03b9 \u03b3\u03c1\u03ae\u03b3\u03bf\u03c1\u03b7 \u03b5\u03be\u03b5\u03c1\u03b5\u03cd\u03bd\u03b7\u03c3\u03b7 "
   "\u03c4\u03bf\u03c5 loss landscape, \u03ba\u03b1\u03b9 \u03bf cosine scheduler \u03bc\u03b5\u03b9\u03ce\u03bd\u03b5\u03b9 \u03c3\u03c4\u03b1\u03b4\u03b9\u03b1\u03ba\u03ac \u03c4\u03bf LR \u03b3\u03b9\u03b1 fine-tuning. "
   "\u03a3\u03c4\u03bf CIFAR-100 \u03bc\u03b5 \u03bc\u03b9\u03ba\u03c1\u03cc CNN, \u03c4\u03bf \u03c5\u03c8\u03b7\u03bb\u03cc lr \u03b2\u03bf\u03b7\u03b8\u03ac \u03c4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03bd\u03b1 \u03be\u03b5\u03c6\u03cd\u03b3\u03b5\u03b9 \u03b1\u03c0\u03cc \u03c4\u03bf\u03c0\u03b9\u03ba\u03ac \u03b5\u03bb\u03ac\u03c7\u03b9\u03c3\u03c4\u03b1.")
bl(doc, "lr=0.01 (50.15%): \u039a\u03b1\u03bb\u03ae \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03b7 \u03b1\u03bb\u03bb\u03ac 6% \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03b1. \u03a4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03c3\u03c5\u03b3\u03ba\u03bb\u03af\u03bd\u03b5\u03b9 "
   "\u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03ac \u03b1\u03bb\u03bb\u03ac \u03b4\u03b5\u03bd \u03bc\u03c0\u03bf\u03c1\u03b5\u03af \u03bd\u03b1 \u03b5\u03be\u03b5\u03c1\u03b5\u03c5\u03bd\u03ae\u03c3\u03b5\u03b9 \u03b1\u03c1\u03ba\u03b5\u03c4\u03ac \u03c4\u03bf loss landscape \u03c3\u03c4\u03b1 \u03c0\u03c1\u03ce\u03c4\u03b1 epochs.")
bl(doc, "lr=0.001 (13.73%): \u03a0\u03bf\u03bb\u03cd \u03b1\u03c1\u03b3\u03ae \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03b7. \u03a3\u03b5 30 epochs \u03c4\u03bf loss \u03ad\u03c0\u03b5\u03c3\u03b5 \u03bc\u03cc\u03bb\u03b9\u03c2 \u03b1\u03c0\u03cc 4.60 \u03c3\u03b5 3.80. "
   "\u0398\u03b1 \u03c7\u03c1\u03b5\u03b9\u03b1\u03b6\u03cc\u03c4\u03b1\u03bd 100+ epochs \u03b3\u03b9\u03b1 \u03bd\u03b1 \u03c6\u03c4\u03ac\u03c3\u03b5\u03b9 \u03c4\u03bf \u03b5\u03c0\u03af\u03c0\u03b5\u03b4\u03bf \u03c4\u03bf\u03c5 lr=0.01.")

im(doc, os.path.join(EX1, "report_06_lr_vs_accuracy.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 1.5: Learning rate vs test accuracy")

# ── 1.5 Time efficiency ──────────────────────────────────────────────────
hd(doc, "1.5 \u0391\u03c0\u03bf\u03b4\u03bf\u03c4\u03b9\u03ba\u03cc\u03c4\u03b7\u03c4\u03b1 \u03a7\u03c1\u03cc\u03bd\u03bf\u03c5", 2)
im(doc, os.path.join(EX1, "report_05_time_vs_accuracy.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 1.6: \u03a7\u03c1\u03cc\u03bd\u03bf\u03c2 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7\u03c2 vs test accuracy")
doc.add_paragraph(
    "\u038c\u03bb\u03b1 \u03c4\u03b1 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03b4\u03b9\u03ae\u03c1\u03ba\u03b5\u03c3\u03b1\u03bd 6-7 \u03bb\u03b5\u03c0\u03c4\u03ac, \u03b5\u03ba\u03c4\u03cc\u03c2 \u03b1\u03c0\u03cc \u03c4\u03bf\u03bd RMSprop (~2.8 min) "
    "\u03c0\u03bf\u03c5 \u03ad\u03ba\u03b1\u03bd\u03b5 early stop. \u039f \u03c7\u03c1\u03cc\u03bd\u03bf\u03c2 \u03b4\u03b5\u03bd \u03b4\u03b9\u03b1\u03c6\u03ad\u03c1\u03b5\u03b9 \u03c3\u03b7\u03bc\u03b1\u03bd\u03c4\u03b9\u03ba\u03ac \u03bc\u03b5\u03c4\u03b1\u03be\u03cd optimizers, "
    "\u03ac\u03c1\u03b1 \u03b7 \u03b5\u03c0\u03b9\u03bb\u03bf\u03b3\u03ae \u03b2\u03b1\u03c3\u03af\u03b6\u03b5\u03c4\u03b1\u03b9 \u03ba\u03b1\u03b8\u03b1\u03c1\u03ac \u03c3\u03c4\u03b7\u03bd \u03b1\u03ba\u03c1\u03af\u03b2\u03b5\u03b9\u03b1.")

# ── 1.6 Sample predictions ───────────────────────────────────────────────
hd(doc, "1.6 \u0395\u03bd\u03b4\u03b5\u03b9\u03ba\u03c4\u03b9\u03ba\u03ac \u0391\u03c0\u03bf\u03c4\u03b5\u03bb\u03ad\u03c3\u03bc\u03b1\u03c4\u03b1 \u03a4\u03b1\u03be\u03b9\u03bd\u03cc\u03bc\u03b7\u03c3\u03b7\u03c2", 2)
im(doc, os.path.join(EX1, "LearningRate_0.1", "sample_predictions.png"), Inches(5.0),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 1.7: \u0394\u03b5\u03af\u03b3\u03bc\u03b1\u03c4\u03b1 \u03c0\u03c1\u03bf\u03b2\u03bb\u03ad\u03c8\u03b5\u03c9\u03bd \u2014 \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf (SGD, lr=0.1)")

# ── 1.7 Conclusions ──────────────────────────────────────────────────────
hd(doc, "1.7 \u03a3\u03c5\u03bc\u03c0\u03b5\u03c1\u03ac\u03c3\u03bc\u03b1\u03c4\u03b1 \u0386\u03c3\u03ba\u03b7\u03c3\u03b7\u03c2 1", 2)
bl(doc, "\u039f SGD \u03b5\u03af\u03bd\u03b1\u03b9 \u03bf \u03ba\u03b1\u03c4\u03b1\u03bb\u03bb\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf\u03c2 optimizer \u03b3\u03b9\u03b1 \u03bc\u03b9\u03ba\u03c1\u03ac CNNs, \u03b4\u03b9\u03cc\u03c4\u03b9 \u03b4\u03b5\u03bd \u03c0\u03c1\u03bf\u03c3\u03b1\u03c1\u03bc\u03cc\u03b6\u03b5\u03b9 "
   "\u03c4\u03bf LR \u03b1\u03c5\u03c4\u03cc\u03bc\u03b1\u03c4\u03b1 \u03ba\u03b1\u03b9 \u03b5\u03c0\u03c9\u03c6\u03b5\u03bb\u03b5\u03af\u03c4\u03b1\u03b9 \u03b1\u03c0\u03cc \u03c4\u03bf\u03bd cosine scheduler.")
bl(doc, "\u039f\u03b9 adaptive optimizers (Adam, AdamW, RMSprop) \u03b1\u03c0\u03b1\u03b9\u03c4\u03bf\u03cd\u03bd \u03c0\u03bf\u03bb\u03cd \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf lr (\u2264 0.001) "
   "\u03b3\u03b9\u03b1 \u03bd\u03b1 \u03bc\u03b7\u03bd \u03b1\u03c0\u03bf\u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03bf\u03c0\u03bf\u03b9\u03ae\u03c3\u03bf\u03c5\u03bd \u03c4\u03b7\u03bd \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7.")
bl(doc, "\u03a4\u03bf learning rate \u03b5\u03af\u03bd\u03b1\u03b9 \u03b7 \u03c0\u03b9\u03bf \u03ba\u03c1\u03af\u03c3\u03b9\u03bc\u03b7 \u03c5\u03c0\u03b5\u03c1-\u03c0\u03b1\u03c1\u03ac\u03bc\u03b5\u03c4\u03c1\u03bf\u03c2: \u03b1\u03c0\u03cc 1% (lr=0.01+Adam) "
   "\u03ad\u03c9\u03c2 56.43% (lr=0.1+SGD), \u03b4\u03b7\u03bb\u03b1\u03b4\u03ae \u03b4\u03b9\u03b1\u03c6\u03bf\u03c1\u03ac 55 \u03bc\u03bf\u03bd\u03ac\u03b4\u03c9\u03bd.")
bl(doc, "\u03a4\u03bf 56.43% \u03b5\u03af\u03bd\u03b1\u03b9 \u03b1\u03be\u03b9\u03bf\u03c0\u03c1\u03b5\u03c0\u03ad\u03c2 \u03b3\u03b9\u03b1 \u03ad\u03bd\u03b1 vanilla CNN \u03c7\u03c9\u03c1\u03af\u03c2 data augmentation, "
   "residual connections \u03ae pretrained weights \u03c3\u03b5 100 \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# EXERCISE 2
# ══════════════════════════════════════════════════════════════════════════
hd(doc, "\u0386\u03c3\u03ba\u03b7\u03c3\u03b7 #2: \u039c\u03b5\u03c4\u03b1\u03c6\u03bf\u03c1\u03ac \u039c\u03ac\u03b8\u03b7\u03c3\u03b7\u03c2 (Transfer Learning)")

hd(doc, "2.1 \u0395\u03b9\u03c3\u03b1\u03b3\u03c9\u03b3\u03ae", 2)
doc.add_paragraph(
    "\u0395\u03c6\u03b1\u03c1\u03bc\u03cc\u03c3\u03c4\u03b7\u03ba\u03b5 transfer learning \u03c3\u03c4\u03bf Oxford-IIIT Pet dataset (37 \u03c1\u03ac\u03c4\u03c3\u03b5\u03c2 \u03c3\u03ba\u03cd\u03bb\u03c9\u03bd/\u03b3\u03b1\u03c4\u03ce\u03bd). "
    "\u03a7\u03c1\u03b7\u03c3\u03b9\u03bc\u03bf\u03c0\u03bf\u03b9\u03ae\u03b8\u03b7\u03ba\u03b1\u03bd \u03c0\u03c1\u03bf\u03b5\u03ba\u03c0\u03b1\u03b9\u03b4\u03b5\u03c5\u03bc\u03ad\u03bd\u03b1 \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03b1 ImageNet: \u03c4\u03b1 \u03c3\u03c5\u03bd\u03b5\u03bb\u03b9\u03ba\u03c4\u03b9\u03ba\u03ac \u03c3\u03c4\u03c1\u03ce\u03bc\u03b1\u03c4\u03b1 "
    "(feature extractor) \u03c0\u03b1\u03b3\u03ce\u03bd\u03bf\u03c5\u03bd \u03ba\u03b1\u03b9 \u03b1\u03bd\u03c4\u03b9\u03ba\u03b1\u03b8\u03af\u03c3\u03c4\u03b1\u03c4\u03b1\u03b9 \u03bc\u03cc\u03bd\u03bf \u03c4\u03bf classifier head \u03b3\u03b9\u03b1 37 \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2. "
    "9 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03c3\u03b5 3 \u03c3\u03b5\u03b9\u03c1\u03ad\u03c2 (20 epochs, step scheduler, Adam):")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 1 \u2014 \u0391\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ad\u03c2: ResNet18, ResNet50, AlexNet, VGG16 (\u03cc\u03bb\u03b5\u03c2 frozen)")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 2 \u2014 Frozen vs Fine-tuned: ResNet18 \u03bc\u03b5 \u03c0\u03b1\u03b3\u03c9\u03bc\u03ad\u03bd\u03b1 \u03ae \u03b5\u03bb\u03b5\u03cd\u03b8\u03b5\u03c1\u03b1 \u03b2\u03ac\u03c1\u03b7")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 3 \u2014 Learning rates: 0.0001, 0.001, 0.01 (ResNet18, frozen)")

hd(doc, "2.2 \u03a3\u03c5\u03bd\u03bf\u03c0\u03c4\u03b9\u03ba\u03cc\u03c2 \u03a0\u03af\u03bd\u03b1\u03ba\u03b1\u03c2 \u0391\u03c0\u03bf\u03c4\u03b5\u03bb\u03b5\u03c3\u03bc\u03ac\u03c4\u03c9\u03bd", 2)
cols2 = ["\u03a0\u03b5\u03af\u03c1\u03b1\u03bc\u03b1", "Model", "LR", "Frozen", "Trainable", "Train(%)", "Test(%)", "\u03a7\u03c1\u03cc\u03bd\u03bf\u03c2(min)"]
rows2 = []
for exp in ex2.values():
    hp = exp.get("hyperparameters", exp.get("config", {}))
    mi = exp.get("model_info", {})
    tp = mi.get("trainable_params", "")
    rows2.append([exp["name"], mi.get("name", hp.get("model_name","")),
                  hp.get("learning_rate",""), str(hp.get("freeze_features","")),
                  f'{tp:,}' if isinstance(tp, int) else str(tp),
                  f'{exp["train_acc"]:.1f}', f'{exp["test_acc"]:.2f}',
                  f'{exp["total_time"]/60:.1f}'])
rows2.sort(key=lambda r: float(r[6]), reverse=True)
tbl(doc, cols2, rows2)
doc.add_paragraph("")

im(doc, os.path.join(EX2, "report_01_accuracy_overview.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 2.1: \u0395\u03c0\u03b9\u03c3\u03ba\u03cc\u03c0\u03b7\u03c3\u03b7 \u03b1\u03ba\u03c1\u03af\u03b2\u03b5\u03b9\u03b1\u03c2 \u03cc\u03bb\u03c9\u03bd \u03c4\u03c9\u03bd \u03c0\u03b5\u03b9\u03c1\u03b1\u03bc\u03ac\u03c4\u03c9\u03bd")

# ── 2.3 Architecture comparison ──────────────────────────────────────────
hd(doc, "2.3 \u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 \u0391\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ce\u03bd", 2)
doc.add_paragraph(
    "\u038c\u03bb\u03b5\u03c2 \u03bf\u03b9 \u03b1\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ad\u03c2 \u03c7\u03c1\u03b7\u03c3\u03b9\u03bc\u03bf\u03c0\u03bf\u03b9\u03bf\u03cd\u03bd pretrained ImageNet \u03b2\u03ac\u03c1\u03b7 \u03bc\u03b5 frozen features. "
    "\u039c\u03cc\u03bd\u03bf \u03c4\u03bf classifier head \u03b5\u03ba\u03c0\u03b1\u03b9\u03b4\u03b5\u03cd\u03b5\u03c4\u03b1\u03b9:")

im(doc, os.path.join(EX2, "report_02_architecture_train_curves.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 2.2: \u039a\u03b1\u03bc\u03c0\u03cd\u03bb\u03b5\u03c2 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7\u03c2 \u03b1\u03bd\u03ac \u03b1\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ae")

doc.add_paragraph("\u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7:")
bl(doc, "ResNet50 (90.87%): \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03b1\u03c0\u03bf\u03c4\u03ad\u03bb\u03b5\u03c3\u03bc\u03b1. \u03a4\u03bf ResNet50 \u03ad\u03c7\u03b5\u03b9 50 \u03c3\u03c4\u03c1\u03ce\u03bc\u03b1\u03c4\u03b1 \u03ba\u03b1\u03b9 residual connections, "
   "\u03b5\u03be\u03ac\u03b3\u03bf\u03bd\u03c4\u03b1\u03c2 \u03c0\u03b9\u03bf \u03c0\u03bb\u03bf\u03cd\u03c3\u03b9\u03b1 \u03ba\u03b1\u03b9 \u03b9\u03b5\u03c1\u03b1\u03c1\u03c7\u03b9\u03ba\u03ac features \u03b1\u03c0\u03cc \u03c4\u03bf ImageNet. "
   "\u03a4\u03bf \u03b2\u03ac\u03b8\u03bf\u03c2 \u03b5\u03c0\u03b9\u03c4\u03c1\u03ad\u03c0\u03b5\u03b9 \u03b1\u03bd\u03af\u03c7\u03bd\u03b5\u03c5\u03c3\u03b7 \u03bb\u03b5\u03c0\u03c4\u03ce\u03bd \u03b4\u03b9\u03b1\u03c6\u03bf\u03c1\u03ce\u03bd \u03bc\u03b5\u03c4\u03b1\u03be\u03cd \u03c1\u03b1\u03c4\u03c3\u03ce\u03bd.")
bl(doc, "VGG16 (89.40%): \u03a0\u03bf\u03bb\u03cd \u03ba\u03bf\u03bd\u03c4\u03ac \u03c3\u03c4\u03bf ResNet50 \u03c0\u03b1\u03c1\u03ac \u03c4\u03b1 134M parameters. "
   "\u03a4\u03bf VGG16 \u03ad\u03c7\u03b5\u03b9 \u03b1\u03c0\u03bb\u03ae \u03b1\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ae (3x3 convolutions \u03bc\u03cc\u03bd\u03bf) \u03b1\u03bb\u03bb\u03ac \u03c4\u03b1 \u03c0\u03bf\u03bb\u03bb\u03ac "
   "\u03c3\u03c4\u03c1\u03ce\u03bc\u03b1\u03c4\u03b1 \u03b4\u03af\u03bd\u03bf\u03c5\u03bd \u03b9\u03c3\u03c7\u03c5\u03c1\u03ac features. \u038c\u03bc\u03c9\u03c2 \u03c7\u03c1\u03b5\u03b9\u03ac\u03b6\u03b5\u03c4\u03b1\u03b9 6x \u03c0\u03b5\u03c1\u03b9\u03c3\u03c3\u03cc\u03c4\u03b5\u03c1\u03bf \u03c7\u03c1\u03cc\u03bd\u03bf (37.7 vs 10.5 min).")
bl(doc, "ResNet18 (87.76%): \u039a\u03b1\u03bb\u03ae \u03b9\u03c3\u03bf\u03c1\u03c1\u03bf\u03c0\u03af\u03b1 \u03b1\u03ba\u03c1\u03af\u03b2\u03b5\u03b9\u03b1\u03c2/\u03c4\u03b1\u03c7\u03cd\u03c4\u03b7\u03c4\u03b1\u03c2. \u039c\u03b5 \u03bc\u03cc\u03bd\u03bf 18.981 \u03b5\u03ba\u03c0\u03b1\u03b9\u03b4\u03b5\u03cd\u03c3\u03b9\u03bc\u03b5\u03c2 "
   "\u03c0\u03b1\u03c1\u03b1\u03bc\u03ad\u03c4\u03c1\u03bf\u03c5\u03c2, \u03c0\u03b5\u03c4\u03c5\u03c7\u03b1\u03af\u03bd\u03b5\u03b9 87.76% \u03c3\u03b5 10.5 \u03bb\u03b5\u03c0\u03c4\u03ac \u2014 \u03bc\u03cc\u03bb\u03b9\u03c2 3% \u03ba\u03ac\u03c4\u03c9 \u03b1\u03c0\u03cc \u03c4\u03bf ResNet50.")
bl(doc, "AlexNet (73.43%): \u03a7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf \u03ba\u03b1\u03c4\u03ac 17%. \u03a4\u03bf AlexNet \u03ad\u03c7\u03b5\u03b9 \u03bc\u03cc\u03bd\u03bf 5 conv layers \u03bc\u03b5 \u03bc\u03b5\u03b3\u03ac\u03bb\u03b1 kernels "
   "(11x11, 5x5), \u03c0\u03bf\u03c5 \u03c7\u03ac\u03bd\u03bf\u03c5\u03bd \u03bb\u03b5\u03c0\u03c4\u03bf\u03bc\u03ad\u03c1\u03b5\u03b9\u03b5\u03c2 texture. \u03a0\u03b1\u03c1\u03cc\u03c4\u03b9 \u03c0\u03b5\u03c4\u03c5\u03c7\u03b1\u03af\u03bd\u03b5\u03b9 98.94% \u03c3\u03c4\u03bf train "
   "(overfitting), \u03c4\u03b1 features \u03c4\u03bf\u03c5 \u03b4\u03b5\u03bd \u03bc\u03b5\u03c4\u03b1\u03c6\u03ad\u03c1\u03bf\u03bd\u03c4\u03b1\u03b9 \u03ba\u03b1\u03bb\u03ac \u03c3\u03c4\u03bf \u03c0\u03c1\u03cc\u03b2\u03bb\u03b7\u03bc\u03b1 \u03c4\u03c9\u03bd \u03c1\u03b1\u03c4\u03c3\u03ce\u03bd.")

# ── 2.4 Frozen vs Fine-tuned ─────────────────────────────────────────────
hd(doc, "2.4 Frozen vs Fine-tuned", 2)
doc.add_paragraph(
    "\u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 ResNet18 \u03bc\u03b5 \u03c0\u03b1\u03b3\u03c9\u03bc\u03ad\u03bd\u03b1 (frozen) \u03ae \u03b5\u03bb\u03b5\u03cd\u03b8\u03b5\u03c1\u03b1 (fine-tuned) \u03b2\u03ac\u03c1\u03b7 "
    "\u03c3\u03c4\u03bf feature extractor:")

im(doc, os.path.join(EX2, "report_03_frozen_vs_finetuned.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 2.3: Frozen vs Fine-tuned \u2014 \u039a\u03b1\u03bc\u03c0\u03cd\u03bb\u03b5\u03c2 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7\u03c2")

doc.add_paragraph("\u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7:")
bl(doc, "Frozen (88.12%): \u0395\u03ba\u03c0\u03b1\u03b9\u03b4\u03b5\u03cd\u03bf\u03bd\u03c4\u03b1\u03b9 \u03bc\u03cc\u03bd\u03bf 18.981 \u03c0\u03b1\u03c1\u03ac\u03bc\u03b5\u03c4\u03c1\u03bf\u03b9 (classifier). "
   "\u03a4\u03b1 ImageNet features \u03b5\u03af\u03bd\u03b1\u03b9 \u03ae\u03b4\u03b7 \u03b5\u03be\u03b1\u03b9\u03c1\u03b5\u03c4\u03b9\u03ba\u03ac \u03b3\u03b5\u03bd\u03b9\u03ba\u03ac \u03ba\u03b1\u03b9 \u03bc\u03b5\u03c4\u03b1\u03c6\u03ad\u03c1\u03bf\u03bd\u03c4\u03b1\u03b9 \u03ac\u03bc\u03b5\u03c3\u03b1. "
   "\u03a4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03b4\u03b5\u03bd \u03ba\u03ac\u03bd\u03b5\u03b9 overfit (train 95.76%) \u03ba\u03b1\u03b9 \u03b3\u03b5\u03bd\u03b9\u03ba\u03b5\u03cd\u03b5\u03b9 \u03ba\u03b1\u03bb\u03ac.")
bl(doc, "Fine-tuned (79.83%): \u0395\u03ba\u03c0\u03b1\u03b9\u03b4\u03b5\u03cd\u03bf\u03bd\u03c4\u03b1\u03b9 \u03cc\u03bb\u03b5\u03c2 \u03bf\u03b9 11.2M \u03c0\u03b1\u03c1\u03ac\u03bc\u03b5\u03c4\u03c1\u03bf\u03b9. "
   "\u03a4\u03bf train accuracy \u03c6\u03c4\u03ac\u03bd\u03b5\u03b9 99.97%, \u03b1\u03bb\u03bb\u03ac \u03c4\u03bf test \u03c0\u03ad\u03c6\u03c4\u03b5\u03b9 \u03c3\u03c4\u03bf 79.83% \u2014 "
   "\u03c3\u03bf\u03b2\u03b1\u03c1\u03cc overfitting. \u039c\u03b5 \u03bc\u03cc\u03bd\u03bf ~3.680 \u03b5\u03b9\u03ba\u03cc\u03bd\u03b5\u03c2 train, \u03c4\u03bf \u03b4\u03af\u03ba\u03c4\u03c5\u03bf \u03bc\u03b5 11M "
   "\u03c0\u03b1\u03c1\u03b1\u03bc\u03ad\u03c4\u03c1\u03bf\u03c5\u03c2 \u03b1\u03c0\u03bf\u03bc\u03bd\u03b7\u03bc\u03bf\u03bd\u03b5\u03cd\u03b5\u03b9 \u03c4\u03b1 train data \u03b1\u03bd\u03c4\u03af \u03bd\u03b1 \u03b3\u03b5\u03bd\u03b9\u03ba\u03b5\u03cd\u03b5\u03b9.")
bl(doc, "\u03a3\u03c5\u03bc\u03c0\u03ad\u03c1\u03b1\u03c3\u03bc\u03b1: \u039c\u03b5 \u03bc\u03b9\u03ba\u03c1\u03cc dataset, \u03c4\u03bf freezing \u03b5\u03af\u03bd\u03b1\u03b9 \u03ba\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03b3\u03b9\u03b1\u03c4\u03af \u03bb\u03b5\u03b9\u03c4\u03bf\u03c5\u03c1\u03b3\u03b5\u03af \u03c9\u03c2 regularization. "
   "\u03a4\u03bf fine-tuning \u03c7\u03c1\u03b5\u03b9\u03ac\u03b6\u03b5\u03c4\u03b1\u03b9 \u03c0\u03bf\u03bb\u03cd \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc lr, data augmentation, \u03ae dropout \u03b3\u03b9\u03b1 \u03bd\u03b1 \u03b1\u03c0\u03bf\u03c6\u03b5\u03c5\u03c7\u03b8\u03b5\u03af \u03c4\u03bf overfitting.")

# ── 2.5 Learning Rate comparison ─────────────────────────────────────────
hd(doc, "2.5 \u0395\u03c0\u03af\u03b4\u03c1\u03b1\u03c3\u03b7 Learning Rate", 2)
doc.add_paragraph(
    "\u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 \u03c4\u03c1\u03b9\u03ce\u03bd learning rates \u03bc\u03b5 ResNet18 frozen:")

im(doc, os.path.join(EX2, "report_04_lr_train_loss_curves.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 2.4: \u039a\u03b1\u03bc\u03c0\u03cd\u03bb\u03b5\u03c2 training loss \u03b1\u03bd\u03ac learning rate")

doc.add_paragraph("\u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7:")
bl(doc, "lr=0.001 (87.71%): \u0399\u03b4\u03b1\u03bd\u03b9\u03ba\u03cc. \u039f Adam \u03bb\u03b5\u03b9\u03c4\u03bf\u03c5\u03c1\u03b3\u03b5\u03af \u03ac\u03c1\u03b9\u03c3\u03c4\u03b1 \u03bc\u03b5 \u03b1\u03c5\u03c4\u03cc \u03c4\u03bf lr \u03c3\u03c4\u03bf "
   "transfer learning, \u03b3\u03b9\u03b1\u03c4\u03af \u03b5\u03af\u03bd\u03b1\u03b9 \u03b1\u03c1\u03ba\u03b5\u03c4\u03ac \u03c5\u03c8\u03b7\u03bb\u03cc \u03b3\u03b9\u03b1 \u03b3\u03c1\u03ae\u03b3\u03bf\u03c1\u03b7 \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03b7, "
   "\u03b1\u03bb\u03bb\u03ac \u03cc\u03c7\u03b9 \u03c4\u03cc\u03c3\u03bf \u03c5\u03c8\u03b7\u03bb\u03cc \u03ce\u03c3\u03c4\u03b5 \u03bd\u03b1 \u03c0\u03c1\u03bf\u03ba\u03b1\u03bb\u03ad\u03c3\u03b5\u03b9 overfit \u03c3\u03c4\u03bf classifier.")
bl(doc, "lr=0.01 (85.99%): \u0395\u03bb\u03b1\u03c6\u03c1\u03cd overfitting. \u03a4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03c6\u03c4\u03ac\u03bd\u03b5\u03b9 98.83% train \u03b1\u03bb\u03bb\u03ac "
   "86% test. \u03a4\u03bf \u03c5\u03c8\u03b7\u03bb\u03cc lr \u03ba\u03ac\u03bd\u03b5\u03b9 \u03c4\u03bf classifier \u03bd\u03b1 \u03c0\u03c1\u03bf\u03c3\u03b1\u03c1\u03bc\u03cc\u03b6\u03b5\u03c4\u03b1\u03b9 \u03c5\u03c0\u03b5\u03c1\u03b2\u03bf\u03bb\u03b9\u03ba\u03ac \u03c3\u03c4\u03b1 train data. "
   "\u03a0\u03b1\u03c1\u03b1\u03c4\u03b7\u03c1\u03b5\u03af\u03c4\u03b1\u03b9 \u03c3\u03b7\u03bc\u03b1\u03bd\u03c4\u03b9\u03ba\u03ae \u03c0\u03c4\u03ce\u03c3\u03b7 loss \u03c3\u03c4\u03bf epoch 7 (\u03bb\u03cc\u03b3\u03c9 step scheduler).")
bl(doc, "lr=0.0001 (80.43%): \u03a0\u03bf\u03bb\u03cd \u03b1\u03c1\u03b3\u03ae \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03b7. \u03a3\u03b5 20 epochs \u03c4\u03bf train accuracy "
   "\u03ad\u03c6\u03c4\u03b1\u03c3\u03b5 \u03bc\u03cc\u03bb\u03b9\u03c2 84.24%. \u03a4\u03bf loss \u03b4\u03b5\u03bd \u03ba\u03b1\u03c4\u03ad\u03b2\u03b7\u03ba\u03b5 \u03b1\u03c1\u03ba\u03b5\u03c4\u03ac (1.22 vs 0.23 \u03bc\u03b5 lr=0.001). "
   "\u0398\u03b1 \u03c7\u03c1\u03b5\u03b9\u03b1\u03b6\u03cc\u03c4\u03b1\u03bd 60+ epochs \u03b3\u03b9\u03b1 \u03bd\u03b1 \u03c6\u03c4\u03ac\u03c3\u03b5\u03b9 \u03c4\u03bf \u03b5\u03c0\u03af\u03c0\u03b5\u03b4\u03bf \u03c4\u03bf\u03c5 lr=0.001.")

im(doc, os.path.join(EX2, "report_06_lr_vs_accuracy.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 2.5: Learning rate vs test accuracy")

# ── 2.6 Time efficiency ──────────────────────────────────────────────────
hd(doc, "2.6 \u0391\u03c0\u03bf\u03b4\u03bf\u03c4\u03b9\u03ba\u03cc\u03c4\u03b7\u03c4\u03b1 \u03a7\u03c1\u03cc\u03bd\u03bf\u03c5", 2)
im(doc, os.path.join(EX2, "report_05_time_vs_accuracy.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 2.6: \u03a7\u03c1\u03cc\u03bd\u03bf\u03c2 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7\u03c2 vs test accuracy")
doc.add_paragraph(
    "\u039f \u03c7\u03c1\u03cc\u03bd\u03bf\u03c2 \u03b4\u03b9\u03b1\u03c6\u03ad\u03c1\u03b5\u03b9 \u03ad\u03bd\u03c4\u03bf\u03bd\u03b1: AlexNet (3.9 min) vs VGG16 (37.7 min). "
    "\u03a4\u03bf ResNet18 frozen \u03c0\u03c1\u03bf\u03c3\u03c6\u03ad\u03c1\u03b5\u03b9 \u03c4\u03b7\u03bd \u03ba\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03b7 \u03b9\u03c3\u03bf\u03c1\u03c1\u03bf\u03c0\u03af\u03b1 "
    "\u03b1\u03ba\u03c1\u03af\u03b2\u03b5\u03b9\u03b1\u03c2/\u03c7\u03c1\u03cc\u03bd\u03bf\u03c5 (87.76% \u03c3\u03b5 10.5 min), \u03b5\u03bd\u03ce \u03c4\u03bf fine-tuning \u03b4\u03b9\u03c0\u03bb\u03b1\u03c3\u03b9\u03ac\u03b6\u03b5\u03b9 "
    "\u03c4\u03bf\u03bd \u03c7\u03c1\u03cc\u03bd\u03bf (25.7 min) \u03c7\u03c9\u03c1\u03af\u03c2 \u03cc\u03c6\u03b5\u03bb\u03bf\u03c2 \u03c3\u03c4\u03b7\u03bd \u03b1\u03ba\u03c1\u03af\u03b2\u03b5\u03b9\u03b1.")

# ── 2.7 Sample predictions ───────────────────────────────────────────────
hd(doc, "2.7 \u0395\u03bd\u03b4\u03b5\u03b9\u03ba\u03c4\u03b9\u03ba\u03ac \u0391\u03c0\u03bf\u03c4\u03b5\u03bb\u03ad\u03c3\u03bc\u03b1\u03c4\u03b1 \u03a4\u03b1\u03be\u03b9\u03bd\u03cc\u03bc\u03b7\u03c3\u03b7\u03c2", 2)
im(doc, os.path.join(EX2, "Architecture_resnet50", "sample_predictions.png"), Inches(5.0),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 2.7: \u0394\u03b5\u03af\u03b3\u03bc\u03b1\u03c4\u03b1 \u03c0\u03c1\u03bf\u03b2\u03bb\u03ad\u03c8\u03b5\u03c9\u03bd \u2014 \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf (ResNet50, frozen)")

# ── 2.8 Conclusions ──────────────────────────────────────────────────────
hd(doc, "2.8 \u03a3\u03c5\u03bc\u03c0\u03b5\u03c1\u03ac\u03c3\u03bc\u03b1\u03c4\u03b1 \u0386\u03c3\u03ba\u03b7\u03c3\u03b7\u03c2 2", 2)
bl(doc, "\u03a4\u03bf transfer learning \u03c0\u03b5\u03c4\u03c5\u03c7\u03b1\u03af\u03bd\u03b5\u03b9 90.87% \u03bc\u03b5 \u03bc\u03cc\u03bd\u03bf 75.813 trainable parameters (ResNet50 frozen) \u2014 "
   "\u03b1\u03c0\u03cc 56.43% \u03bc\u03b5 964K parameters \u03c3\u03c4\u03bf vanilla CNN (\u0386\u03c3\u03ba\u03b7\u03c3\u03b7 1).")
bl(doc, "\u039f\u03b9 \u03b2\u03b1\u03b8\u03cd\u03c4\u03b5\u03c1\u03b5\u03c2 \u03b1\u03c1\u03c7\u03b9\u03c4\u03b5\u03ba\u03c4\u03bf\u03bd\u03b9\u03ba\u03ad\u03c2 \u03b5\u03be\u03ac\u03b3\u03bf\u03c5\u03bd \u03ba\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03b1 features: ResNet50 > VGG16 > ResNet18 > AlexNet.")
bl(doc, "\u03a4\u03bf freezing \u03c5\u03c0\u03b5\u03c1\u03c4\u03b5\u03c1\u03b5\u03af \u03c4\u03bf\u03c5 fine-tuning \u03c3\u03b5 \u03bc\u03b9\u03ba\u03c1\u03ac datasets, \u03bb\u03b5\u03b9\u03c4\u03bf\u03c5\u03c1\u03b3\u03ce\u03bd\u03c4\u03b1\u03c2 \u03c9\u03c2 \u03b9\u03c3\u03c7\u03c5\u03c1\u03cc regularizer.")
bl(doc, "\u03a4\u03bf lr=0.001 \u03b5\u03af\u03bd\u03b1\u03b9 \u03b9\u03b4\u03b1\u03bd\u03b9\u03ba\u03cc \u03b3\u03b9\u03b1 Adam+frozen features. \u03a4\u03bf lr \u03b1\u03c6\u03bf\u03c1\u03ac \u03bc\u03cc\u03bd\u03bf \u03c4\u03bf classifier, "
   "\u03ac\u03c1\u03b1 \u03c7\u03c1\u03b5\u03b9\u03ac\u03b6\u03b5\u03c4\u03b1\u03b9 \u03b1\u03c1\u03ba\u03b5\u03c4\u03ae \u03c4\u03b1\u03c7\u03cd\u03c4\u03b7\u03c4\u03b1 \u03c7\u03c9\u03c1\u03af\u03c2 \u03bd\u03b1 \u03c0\u03c1\u03bf\u03ba\u03b1\u03bb\u03b5\u03af overfitting.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# EXERCISE 3
# ══════════════════════════════════════════════════════════════════════════
hd(doc, "\u0386\u03c3\u03ba\u03b7\u03c3\u03b7 #3: \u03a3\u03b7\u03bc\u03b1\u03c3\u03b9\u03bf\u03bb\u03bf\u03b3\u03b9\u03ba\u03ae \u03a4\u03bc\u03b7\u03bc\u03b1\u03c4\u03bf\u03c0\u03bf\u03af\u03b7\u03c3\u03b7 (Semantic Segmentation)")

hd(doc, "3.1 \u0395\u03b9\u03c3\u03b1\u03b3\u03c9\u03b3\u03ae", 2)
doc.add_paragraph(
    "\u03a5\u03bb\u03bf\u03c0\u03bf\u03b9\u03ae\u03b8\u03b7\u03ba\u03b5 U-Net (31M \u03c0\u03b1\u03c1\u03ac\u03bc\u03b5\u03c4\u03c1\u03bf\u03b9, base channels=64) \u03b3\u03b9\u03b1 semantic segmentation "
    "\u03c3\u03c4\u03bf SBD dataset (21 \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2 Pascal VOC). \u03a4\u03bf U-Net \u03c7\u03c1\u03b7\u03c3\u03b9\u03bc\u03bf\u03c0\u03bf\u03b9\u03b5\u03af encoder-decoder "
    "\u03bc\u03b5 skip connections \u03b3\u03b9\u03b1 \u03bd\u03b1 \u03b4\u03b9\u03b1\u03c4\u03b7\u03c1\u03b5\u03af \u03c7\u03c9\u03c1\u03b9\u03ba\u03ad\u03c2 \u03bb\u03b5\u03c0\u03c4\u03bf\u03bc\u03ad\u03c1\u03b5\u03b9\u03b5\u03c2. "
    "7 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03c3\u03b5 2 \u03c3\u03b5\u03b9\u03c1\u03ad\u03c2:")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 1 \u2014 \u039a\u03b1\u03bd\u03bf\u03bd\u03b9\u03ba\u03cc BG weight=1.0: 3 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1, 10 epochs, step scheduler, CrossEntropy")
bl(doc, "\u03a3\u03b5\u03b9\u03c1\u03ac 2 \u2014 \u039c\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03bf BG weight=0.2: 4 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1, 5 epochs, cosine scheduler, WeightedCE")

hd(doc, "3.2 \u03a1\u03cc\u03bb\u03bf\u03c2 Background Weight", 2)
doc.add_paragraph(
    "\u03a3\u03c4\u03bf SBD, \u03c4\u03bf background \u03ba\u03b1\u03c4\u03b1\u03bb\u03b1\u03bc\u03b2\u03ac\u03bd\u03b5\u03b9 >80% \u03c4\u03c9\u03bd pixels \u03c3\u03b5 \u03ba\u03ac\u03b8\u03b5 \u03b5\u03b9\u03ba\u03cc\u03bd\u03b1. "
    "\u039c\u03b5 weight=1.0, \u03c4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03b2\u03b5\u03bb\u03c4\u03b9\u03c3\u03c4\u03bf\u03c0\u03bf\u03b9\u03b5\u03af \u03ba\u03c5\u03c1\u03af\u03c9\u03c2 \u03b3\u03b9\u03b1 \u03c4\u03bf background \u03ba\u03b1\u03b9 "
    "\u03c0\u03b5\u03c4\u03c5\u03c7\u03b1\u03af\u03bd\u03b5\u03b9 \u03c5\u03c8\u03b7\u03bb\u03cc Pixel Accuracy \u03b1\u03bb\u03bb\u03ac \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc mIoU \u03c3\u03c4\u03b9\u03c2 foreground \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2. "
    "\u039c\u03b5\u03b9\u03ce\u03bd\u03bf\u03bd\u03c4\u03b1\u03c2 \u03c4\u03bf \u03c3\u03c4\u03bf 0.2, \u03b7 loss \u03b4\u03af\u03bd\u03b5\u03b9 \u03c0\u03b5\u03c1\u03b9\u03c3\u03c3\u03cc\u03c4\u03b5\u03c1\u03bf \u03b2\u03ac\u03c1\u03bf\u03c2 \u03c3\u03c4\u03b9\u03c2 foreground \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2, "
    "\u03b1\u03bb\u03bb\u03ac \u03b7 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7 \u03b3\u03af\u03bd\u03b5\u03c4\u03b1\u03b9 \u03c0\u03b9\u03bf \u03b1\u03c3\u03c4\u03b1\u03b8\u03ae\u03c2 \u03ba\u03b1\u03b9 \u03c7\u03c1\u03b5\u03b9\u03ac\u03b6\u03b5\u03c4\u03b1\u03b9 \u03c0\u03b5\u03c1\u03b9\u03c3\u03c3\u03cc\u03c4\u03b5\u03c1\u03b1 epochs.")

hd(doc, "3.3 \u03a3\u03c5\u03bd\u03bf\u03c0\u03c4\u03b9\u03ba\u03cc\u03c2 \u03a0\u03af\u03bd\u03b1\u03ba\u03b1\u03c2 \u0391\u03c0\u03bf\u03c4\u03b5\u03bb\u03b5\u03c3\u03bc\u03ac\u03c4\u03c9\u03bd", 2)
cols3 = ["\u03a0\u03b5\u03af\u03c1\u03b1\u03bc\u03b1", "BG Wt", "Optimizer", "LR", "Epochs", "mIoU(%)", "PixAcc(%)", "\u03a7\u03c1\u03cc\u03bd\u03bf\u03c2(min)"]
rows3 = []
for exp in ex3n.values():
    hp = exp.get("hyperparameters", exp.get("config", {}))
    ep = len(exp["history"]["train_loss"])
    rows3.append([exp["name"], "1.0", hp.get("optimizer","").upper(), hp.get("learning_rate",""),
                  ep, f'{exp["val_miou"]:.2f}', f'{exp["val_pixel_acc"]:.2f}', f'{exp["total_time"]/60:.1f}'])
for exp in ex3l.values():
    hp = exp.get("hyperparameters", exp.get("config", {}))
    ep = len(exp["history"]["train_loss"])
    rows3.append([exp["name"], hp.get("background_weight","0.2"), hp.get("optimizer","").upper(),
                  hp.get("learning_rate",""), ep, f'{exp["val_miou"]:.2f}', f'{exp["val_pixel_acc"]:.2f}',
                  f'{exp["total_time"]/60:.1f}'])
rows3.sort(key=lambda r: float(r[5]), reverse=True)
tbl(doc, cols3, rows3)
doc.add_paragraph("")

# ── 3.4 Series 1: Normal BG Weight ──────────────────────────────────────
hd(doc, "3.4 \u03a3\u03b5\u03b9\u03c1\u03ac 1: \u039a\u03b1\u03bd\u03bf\u03bd\u03b9\u03ba\u03cc BG Weight (=1.0)", 2)
doc.add_paragraph(
    "3 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03bc\u03b5 10 epochs, step scheduler. "
    "\u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 AdamW vs SGD \u03ba\u03b1\u03b9 \u03b4\u03b9\u03b1\u03c6\u03bf\u03c1\u03b5\u03c4\u03b9\u03ba\u03ce\u03bd learning rates:")

im(doc, os.path.join(EX3N, "report_01_miou_overview.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 3.1: \u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 mIoU \u2014 \u039a\u03b1\u03bd\u03bf\u03bd\u03b9\u03ba\u03cc BG Weight")

im(doc, os.path.join(EX3N, "report_02_optimizer_miou_curves.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 3.2: \u039a\u03b1\u03bc\u03c0\u03cd\u03bb\u03b5\u03c2 mIoU \u03b1\u03bd\u03ac \u03c0\u03b5\u03af\u03c1\u03b1\u03bc\u03b1 (BG=1.0)")

doc.add_paragraph("\u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7:")
bl(doc, "BaseChannels_64 (AdamW, lr=0.0003): \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03bc\u03b5 mIoU=13.91%, PixAcc=74.79%. "
   "\u039f AdamW \u03c3\u03c5\u03bd\u03b4\u03c5\u03ac\u03b6\u03b5\u03b9 adaptive LR \u03bc\u03b5 weight decay, \u03c0\u03c1\u03bf\u03c3\u03c6\u03ad\u03c1\u03bf\u03bd\u03c4\u03b1\u03c2 "
   "\u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03ae \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03b7 \u03c3\u03b5 \u03ad\u03bd\u03b1 \u03b4\u03af\u03ba\u03c4\u03c5\u03bf \u03bc\u03b5 31M \u03c0\u03b1\u03c1\u03b1\u03bc\u03ad\u03c4\u03c1\u03bf\u03c5\u03c2.")
bl(doc, "Optimizer_SGD (lr=0.0003): \u03a7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf mIoU=10.19%, PixAcc=71.11%. "
   "\u039f SGD \u03c7\u03c9\u03c1\u03af\u03c2 adaptive scaling \u03c3\u03c5\u03b3\u03ba\u03bb\u03af\u03bd\u03b5\u03b9 \u03c0\u03b9\u03bf \u03b1\u03c1\u03b3\u03ac \u03c3\u03b5 \u03bc\u03b5\u03b3\u03ac\u03bb\u03b1 \u03b4\u03af\u03ba\u03c4\u03c5\u03b1. "
   "\u03a3\u03b5 \u03b1\u03bd\u03c4\u03af\u03b8\u03b5\u03c3\u03b7 \u03bc\u03b5 \u03c4\u03b7\u03bd \u0386\u03c3\u03ba\u03b7\u03c3\u03b7 1 (\u03cc\u03c0\u03bf\u03c5 \u03bf SGD \u03ae\u03c4\u03b1\u03bd \u03ba\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf\u03c2), "
   "\u03b5\u03b4\u03ce \u03c4\u03bf \u03bc\u03ad\u03b3\u03b5\u03b8\u03bf\u03c2 \u03c4\u03bf\u03c5 U-Net \u03b1\u03c0\u03b1\u03b9\u03c4\u03b5\u03af adaptive optimizer.")
bl(doc, "LearningRate_0.001 (AdamW): mIoU=12.88%, PixAcc=73.65%. \u03a4\u03bf \u03c5\u03c8\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf LR "
   "\u03c0\u03c1\u03bf\u03ba\u03b1\u03bb\u03b5\u03af \u03b5\u03bb\u03b1\u03c6\u03c1\u03ce\u03c2 \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf mIoU \u2014 \u03c4\u03bf 0.0003 \u03b5\u03af\u03bd\u03b1\u03b9 \u03c0\u03b9\u03bf \u03ba\u03b1\u03c4\u03ac\u03bb\u03bb\u03b7\u03bb\u03bf "
   "\u03b3\u03b9\u03b1 fine-grained segmentation \u03cc\u03c0\u03bf\u03c5 \u03b7 \u03b1\u03ba\u03c1\u03af\u03b2\u03b5\u03b9\u03b1 pixel-level \u03b5\u03af\u03bd\u03b1\u03b9 \u03ba\u03c1\u03af\u03c3\u03b9\u03bc\u03b7.")

im(doc, os.path.join(EX3N, "report_07_best_segmentation_panel.png"), Inches(5.2),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 3.3: \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03b1\u03c0\u03bf\u03c4\u03ad\u03bb\u03b5\u03c3\u03bc\u03b1 \u03c4\u03bc\u03b7\u03bc\u03b1\u03c4\u03bf\u03c0\u03bf\u03af\u03b7\u03c3\u03b7\u03c2 (BG=1.0, AdamW)")

im(doc, os.path.join(EX3N, "report_05_time_vs_miou.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 3.4: \u03a7\u03c1\u03cc\u03bd\u03bf\u03c2 vs mIoU (BG=1.0)")

# ── 3.5 Series 2: Reduced BG Weight ─────────────────────────────────────
hd(doc, "3.5 \u03a3\u03b5\u03b9\u03c1\u03ac 2: \u039c\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03bf BG Weight (=0.2)", 2)
doc.add_paragraph(
    "4 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03bc\u03b5 5 epochs, cosine scheduler. "
    "\u03a3\u03c4\u03cc\u03c7\u03bf\u03c2: \u03bd\u03b1 \u03b4\u03bf\u03cd\u03bc\u03b5 \u03b1\u03bd \u03b7 \u03bc\u03b5\u03af\u03c9\u03c3\u03b7 \u03c4\u03bf\u03c5 BG weight "
    "\u03b2\u03b5\u03bb\u03c4\u03b9\u03ce\u03bd\u03b5\u03b9 \u03c4\u03bf foreground mIoU:")

im(doc, os.path.join(EX3L, "report_01_miou_overview.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 3.5: \u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 mIoU \u2014 \u039c\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03bf BG Weight (0.2)")

im(doc, os.path.join(EX3L, "report_02_optimizer_miou_curves.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 3.6: \u039a\u03b1\u03bc\u03c0\u03cd\u03bb\u03b5\u03c2 mIoU \u03b1\u03bd\u03ac \u03c0\u03b5\u03af\u03c1\u03b1\u03bc\u03b1 (BG=0.2)")

doc.add_paragraph("\u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7:")
bl(doc, "LR_0.0001 + AdamW (mIoU=7.56%): \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03c4\u03b7\u03c2 \u03c3\u03b5\u03b9\u03c1\u03ac\u03c2 2. "
   "\u03a4\u03bf \u03c0\u03bf\u03bb\u03cd \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc lr \u03c3\u03c4\u03b1\u03b8\u03b5\u03c1\u03bf\u03c0\u03bf\u03b9\u03b5\u03af \u03c4\u03b7\u03bd \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7 \u03c0\u03b1\u03c1\u03ac "
   "\u03c4\u03b1 \u03b1\u03bd\u03b9\u03c3\u03cc\u03c1\u03c1\u03bf\u03c0\u03b1 class weights. \u039c\u03b5 \u03c0\u03b5\u03c1\u03b9\u03c3\u03c3\u03cc\u03c4\u03b5\u03c1\u03b1 epochs \u03b1\u03bd\u03b1\u03bc\u03ad\u03bd\u03b5\u03c4\u03b1\u03b9 \u03c3\u03b7\u03bc\u03b1\u03bd\u03c4\u03b9\u03ba\u03ae \u03b2\u03b5\u03bb\u03c4\u03af\u03c9\u03c3\u03b7.")
bl(doc, "LR_0.001 + AdamW (mIoU=5.55%): \u03a4\u03bf \u03c5\u03c8\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf LR \u03c3\u03b5 \u03c3\u03c5\u03bd\u03b4\u03c5\u03b1\u03c3\u03bc\u03cc "
   "\u03bc\u03b5 BG=0.2 \u03c0\u03c1\u03bf\u03ba\u03b1\u03bb\u03b5\u03af \u03b1\u03c3\u03c4\u03ac\u03b8\u03b5\u03b9\u03b1: \u03c4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03c5\u03c0\u03b5\u03c1-\u03b1\u03bd\u03c4\u03b9\u03b4\u03c1\u03ac \u03c3\u03c4\u03b1 "
   "\u03bc\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03b1 background gradients \u03ba\u03b1\u03b9 \u03b4\u03b5\u03bd \u03c3\u03c5\u03b3\u03ba\u03bb\u03af\u03bd\u03b5\u03b9 \u03c3\u03c9\u03c3\u03c4\u03ac.")
bl(doc, "SGD (mIoU=3.04-5.07%): \u0391\u03ba\u03cc\u03bc\u03b1 \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03b1. \u039f SGD \u03b4\u03b5\u03bd \u03bc\u03c0\u03bf\u03c1\u03b5\u03af \u03bd\u03b1 \u03c0\u03c1\u03bf\u03c3\u03b1\u03c1\u03bc\u03bf\u03c3\u03c4\u03b5\u03af "
   "\u03c3\u03c4\u03b1 \u03b1\u03bd\u03b9\u03c3\u03cc\u03c1\u03c1\u03bf\u03c0\u03b1 gradients \u03c0\u03bf\u03c5 \u03c0\u03c1\u03bf\u03ba\u03b1\u03bb\u03b5\u03af \u03c4\u03bf BG=0.2 \u03c3\u03b5 \u03bc\u03cc\u03bb\u03b9\u03c2 5 epochs.")

im(doc, os.path.join(EX3L, "report_04_lr_val_loss_curves.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 3.7: Validation loss \u03b1\u03bd\u03ac LR (BG=0.2)")

im(doc, os.path.join(EX3L, "report_07_best_segmentation_panel.png"), Inches(5.2),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 3.8: \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03b1\u03c0\u03bf\u03c4\u03ad\u03bb\u03b5\u03c3\u03bc\u03b1 \u03c4\u03bc\u03b7\u03bc\u03b1\u03c4\u03bf\u03c0\u03bf\u03af\u03b7\u03c3\u03b7\u03c2 (BG=0.2)")

# ── 3.6 Comparison BG=1.0 vs BG=0.2 ─────────────────────────────────────
hd(doc, "3.6 \u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 BG=1.0 vs BG=0.2", 2)
doc.add_paragraph("\u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7:")
bl(doc, "BG=1.0: \u03a3\u03c4\u03b1\u03b8\u03b5\u03c1\u03ae \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7 \u03ba\u03b1\u03b9 \u03c5\u03c8\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf mIoU (10-14%) \u03ba\u03b1\u03b9 PixAcc (71-75%). "
   "\u03a4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03bc\u03b1\u03b8\u03b1\u03af\u03bd\u03b5\u03b9 \u03ba\u03b1\u03bb\u03ac \u03c4\u03bf background \u03ba\u03b1\u03b9 \u03c3\u03c4\u03b1\u03b4\u03b9\u03b1\u03ba\u03ac \u03b2\u03b5\u03bb\u03c4\u03b9\u03ce\u03bd\u03b5\u03b9 "
   "\u03ba\u03b1\u03b9 \u03c4\u03b9\u03c2 foreground \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2.")
bl(doc, "BG=0.2: \u03a7\u03b1\u03bc\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03b1 mIoU (3-7.5%) \u03ba\u03b1\u03b9 PixAcc (13-55%). "
   "\u0397 \u03ba\u03cd\u03c1\u03b9\u03b1 \u03b1\u03b9\u03c4\u03af\u03b1 \u03b5\u03af\u03bd\u03b1\u03b9 \u03cc\u03c4\u03b9 5 epochs \u03b4\u03b5\u03bd \u03b5\u03af\u03bd\u03b1\u03b9 \u03b1\u03c1\u03ba\u03b5\u03c4\u03ac \u2014 "
   "\u03c4\u03bf \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf \u03b4\u03b5\u03bd \u03c0\u03c1\u03bf\u03bb\u03b1\u03b2\u03b1\u03af\u03bd\u03b5\u03b9 \u03bd\u03b1 \u03c3\u03c5\u03b3\u03ba\u03bb\u03af\u03bd\u03b5\u03b9. "
   "\u0395\u03c0\u03af\u03c3\u03b7\u03c2, \u03c4\u03bf \u03bc\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03bf BG weight \u03c0\u03c1\u03bf\u03ba\u03b1\u03bb\u03b5\u03af \u03bc\u03b5\u03b3\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03b7 gradient variance.")
bl(doc, "\u039c\u03b5 \u03af\u03c3\u03b1 epochs (10), \u03c4\u03bf \u03bc\u03b5\u03b9\u03c9\u03bc\u03ad\u03bd\u03bf BG weight \u03b1\u03bd\u03b1\u03bc\u03ad\u03bd\u03b5\u03c4\u03b1\u03b9 \u03bd\u03b1 \u03c5\u03c0\u03b5\u03c1\u03c4\u03b5\u03c1\u03ae\u03c3\u03b5\u03b9 "
   "\u03c3\u03c4\u03bf foreground mIoU, \u03b3\u03b9\u03b1\u03c4\u03af \u03b8\u03b1 \u03b4\u03ce\u03c3\u03b5\u03b9 \u03c0\u03b5\u03c1\u03b9\u03c3\u03c3\u03cc\u03c4\u03b5\u03c1\u03b7 \u03c0\u03c1\u03bf\u03c3\u03bf\u03c7\u03ae "
   "\u03c3\u03c4\u03b9\u03c2 \u03c3\u03c0\u03ac\u03bd\u03b9\u03b5\u03c2 \u03ba\u03bb\u03ac\u03c3\u03b5\u03b9\u03c2 (person, car, dog \u03ba\u03bb\u03c0).")

# ── 3.7 Conclusions ──────────────────────────────────────────────────────
hd(doc, "3.7 \u03a3\u03c5\u03bc\u03c0\u03b5\u03c1\u03ac\u03c3\u03bc\u03b1\u03c4\u03b1 \u0386\u03c3\u03ba\u03b7\u03c3\u03b7\u03c2 3", 2)
bl(doc, "\u039f AdamW \u03b5\u03af\u03bd\u03b1\u03b9 \u03bf \u03ba\u03b1\u03c4\u03b1\u03bb\u03bb\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf\u03c2 optimizer \u03b3\u03b9\u03b1 U-Net (31M params), "
   "\u03c3\u03b5 \u03b1\u03bd\u03c4\u03af\u03b8\u03b5\u03c3\u03b7 \u03bc\u03b5 \u03c4\u03bf\u03bd SGD \u03c0\u03bf\u03c5 \u03c5\u03c0\u03b5\u03c1\u03c4\u03b5\u03c1\u03bf\u03cd\u03c3\u03b5 \u03c3\u03c4\u03bf vanilla CNN (964K params).")
bl(doc, "\u03a4\u03bf \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc lr (0.0003) \u03b5\u03af\u03bd\u03b1\u03b9 \u03ba\u03c1\u03af\u03c3\u03b9\u03bc\u03bf \u03b3\u03b9\u03b1 segmentation: "
   "\u03c4\u03bf pixel-level \u03c0\u03c1\u03cc\u03b2\u03bb\u03b7\u03bc\u03b1 \u03b1\u03c0\u03b1\u03b9\u03c4\u03b5\u03af \u03c0\u03bf\u03bb\u03cd \u03b1\u03ba\u03c1\u03b9\u03b2\u03ae \u03c3\u03cd\u03b3\u03ba\u03bb\u03b9\u03c3\u03b7.")
bl(doc, "\u03a4\u03bf BG weight \u03b5\u03af\u03bd\u03b1\u03b9 \u03ba\u03c1\u03af\u03c3\u03b9\u03bc\u03b7 \u03c5\u03c0\u03b5\u03c1-\u03c0\u03b1\u03c1\u03ac\u03bc\u03b5\u03c4\u03c1\u03bf\u03c2: "
   "\u03b7 \u03bc\u03b5\u03af\u03c9\u03c3\u03b7 \u03b1\u03c0\u03cc 1.0 \u03c3\u03b5 0.2 \u03b1\u03bb\u03bb\u03ac\u03b6\u03b5\u03b9 \u03c1\u03b9\u03b6\u03b9\u03ba\u03ac \u03c4\u03b7 \u03c3\u03c5\u03bc\u03c0\u03b5\u03c1\u03b9\u03c6\u03bf\u03c1\u03ac \u03c4\u03bf\u03c5 \u03bc\u03bf\u03bd\u03c4\u03ad\u03bb\u03bf\u03c5 "
   "\u03ba\u03b1\u03b9 \u03b1\u03c0\u03b1\u03b9\u03c4\u03b5\u03af \u03b1\u03bd\u03c4\u03af\u03c3\u03c4\u03bf\u03b9\u03c7\u03b7 \u03c0\u03c1\u03bf\u03c3\u03b1\u03c1\u03bc\u03bf\u03b3\u03ae LR \u03ba\u03b1\u03b9 epochs.")
bl(doc, "\u03a4\u03bf mIoU=13.91% \u03b5\u03af\u03bd\u03b1\u03b9 \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc \u03c3\u03b5 \u03c3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 \u03bc\u03b5 SOTA (\u03c0.\u03c7. DeepLabV3+ >70%), "
   "\u03b1\u03bb\u03bb\u03ac \u03b1\u03bd\u03b1\u03bc\u03b5\u03bd\u03cc\u03bc\u03b5\u03bd\u03bf \u03b3\u03b9\u03b1 U-Net \u03c7\u03c9\u03c1\u03af\u03c2 pretrained encoder \u03c3\u03b5 10 epochs.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# EXERCISE 4
# ══════════════════════════════════════════════════════════════════════════
hd(doc, "\u0386\u03c3\u03ba\u03b7\u03c3\u03b7 #4: \u0391\u03bd\u03af\u03c7\u03bd\u03b5\u03c5\u03c3\u03b7 \u0391\u03bd\u03c4\u03b9\u03ba\u03b5\u03b9\u03bc\u03ad\u03bd\u03c9\u03bd (Object Detection)")

hd(doc, "4.1 \u0395\u03b9\u03c3\u03b1\u03b3\u03c9\u03b3\u03ae", 2)
doc.add_paragraph(
    "\u0395\u03c6\u03b1\u03c1\u03bc\u03cc\u03c3\u03c4\u03b7\u03ba\u03b5 Faster R-CNN \u03c3\u03c4\u03bf Oxford-IIIT Pet dataset "
    "\u03b3\u03b9\u03b1 \u03b5\u03bd\u03c4\u03bf\u03c0\u03b9\u03c3\u03bc\u03cc \u03b6\u03ce\u03c9\u03bd (bounding box + \u03ba\u03bb\u03ac\u03c3\u03b7). "
    "\u03a4\u03bf Faster R-CNN \u03c7\u03c1\u03b7\u03c3\u03b9\u03bc\u03bf\u03c0\u03bf\u03b9\u03b5\u03af Region Proposal Network (RPN) "
    "\u03ba\u03b1\u03b9 pretrained backbone \u03b1\u03c0\u03cc \u03c4\u03bf COCO dataset. "
    "4 \u03c0\u03b5\u03b9\u03c1\u03ac\u03bc\u03b1\u03c4\u03b1 \u03c3\u03b5 2 \u03ac\u03be\u03bf\u03bd\u03b5\u03c2 (4 epochs, SGD, step scheduler):")
bl(doc, "\u0386\u03be\u03bf\u03bd\u03b1\u03c2 1 \u2014 Backbone: ResNet50-FPN vs MobileNetV3 (\u03b2\u03ac\u03b8\u03bf\u03c2 vs \u03c4\u03b1\u03c7\u03cd\u03c4\u03b7\u03c4\u03b1)")
bl(doc, "\u0386\u03be\u03bf\u03bd\u03b1\u03c2 2 \u2014 Learning rate: 0.001 vs 0.005")

hd(doc, "4.2 \u03a3\u03c5\u03bd\u03bf\u03c0\u03c4\u03b9\u03ba\u03cc\u03c2 \u03a0\u03af\u03bd\u03b1\u03ba\u03b1\u03c2 \u0391\u03c0\u03bf\u03c4\u03b5\u03bb\u03b5\u03c3\u03bc\u03ac\u03c4\u03c9\u03bd", 2)
# Load ex4 results from per-experiment JSONs
ex4_data = []
for dn in ["resnet50_lr0.005", "resnet50_lr0.001", "mobilenet_lr0.001", "mobilenet_lr0.005"]:
    jp = os.path.join(EX4, dn, "results.json")
    if os.path.exists(jp):
        with open(jp, encoding="utf-8") as f:
            ex4_data.append(json.load(f))
cols4 = ["\u03a0\u03b5\u03af\u03c1\u03b1\u03bc\u03b1", "Backbone", "LR", "Det Acc(%)", "Train Loss", "Cls Loss", "Box Loss"]
rows4 = []
for e in ex4_data:
    m = e["metrics"]
    c = e["config"]
    rows4.append([e["name"], c["backbone"].upper(), c["learning_rate"],
                  f'{m["detection_accuracy"]:.2f}', f'{m["final_train_loss"]:.4f}',
                  f'{m["final_classifier_loss"]:.4f}', f'{m["final_box_reg_loss"]:.4f}'])
rows4.sort(key=lambda r: float(r[3]), reverse=True)
tbl(doc, cols4, rows4)
doc.add_paragraph("")

im(doc, os.path.join(EX4, "report_01_detection_accuracy_overview.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 4.1: \u0395\u03c0\u03b9\u03c3\u03ba\u03cc\u03c0\u03b7\u03c3\u03b7 \u03b1\u03ba\u03c1\u03af\u03b2\u03b5\u03b9\u03b1\u03c2 \u03b1\u03bd\u03af\u03c7\u03bd\u03b5\u03c5\u03c3\u03b7\u03c2")

# ── 4.3 Backbone comparison ──────────────────────────────────────────────
hd(doc, "4.3 \u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 Backbones", 2)
doc.add_paragraph(
    "\u03a3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 ResNet50-FPN vs MobileNetV3 \u03c9\u03c2 backbone \u03c4\u03bf\u03c5 Faster R-CNN:")

im(doc, os.path.join(EX4, "report_02_total_loss_curves.png"), Inches(4.8),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 4.2: \u039a\u03b1\u03bc\u03c0\u03cd\u03bb\u03b5\u03c2 total loss \u03b1\u03bd\u03ac \u03c0\u03b5\u03af\u03c1\u03b1\u03bc\u03b1")

doc.add_paragraph("\u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7:")
bl(doc, "ResNet50 + lr=0.005 (83.18%): \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf. \u03a4\u03bf ResNet50-FPN \u03c0\u03b1\u03c1\u03ad\u03c7\u03b5\u03b9 "
   "\u03c0\u03bb\u03bf\u03cd\u03c3\u03b9\u03b1 multi-scale features \u03bc\u03ad\u03c3\u03c9 Feature Pyramid Network. "
   "\u039c\u03b5 lr=0.005 \u03c4\u03bf loss \u03ad\u03c0\u03b5\u03c3\u03b5 \u03c3\u03c4\u03bf 0.1998, \u03b5\u03c0\u03b9\u03c4\u03c1\u03ad\u03c0\u03bf\u03bd\u03c4\u03b1\u03c2 "
   "\u03b1\u03ba\u03c1\u03b9\u03b2\u03ae\u03c2 localization \u03ba\u03b1\u03b9 classification.")
bl(doc, "ResNet50 + lr=0.001 (25.65%): \u0391\u03c0\u03bf\u03c4\u03c5\u03c7\u03af\u03b1 \u03c0\u03b1\u03c1\u03ac \u03c4\u03bf \u03af\u03b4\u03b9\u03bf backbone. "
   "\u03a4\u03bf \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc lr \u03b4\u03b5\u03bd \u03b5\u03c0\u03b1\u03c1\u03ba\u03b5\u03af \u03b3\u03b9\u03b1 \u03bd\u03b1 \u03c0\u03c1\u03bf\u03c3\u03b1\u03c1\u03bc\u03cc\u03c3\u03b5\u03b9 "
   "\u03c4\u03bf RPN \u03ba\u03b1\u03b9 \u03c4\u03bf\u03bd detection head \u03c3\u03b5 4 epochs. "
   "\u03a4\u03bf loss \u03ad\u03bc\u03b5\u03b9\u03bd\u03b5 \u03c3\u03c4\u03bf 0.276, \u03c0\u03bf\u03bb\u03cd \u03c5\u03c8\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf \u03b1\u03c0\u03cc \u03c4\u03bf lr=0.005.")
bl(doc, "MobileNet + lr=0.001 (73.78%): \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf \u03b1\u03c0\u03cc \u03c4\u03bf ResNet50 \u03bc\u03b5 \u03af\u03b4\u03b9\u03bf lr! "
   "\u03a4\u03bf MobileNetV3 \u03b5\u03af\u03bd\u03b1\u03b9 \u03b5\u03bb\u03b1\u03c6\u03c1\u03cd\u03c4\u03b5\u03c1\u03bf \u03ba\u03b1\u03b9 \u03c3\u03c5\u03b3\u03ba\u03bb\u03af\u03bd\u03b5\u03b9 "
   "\u03b3\u03c1\u03b7\u03b3\u03bf\u03c1\u03cc\u03c4\u03b5\u03c1\u03b1 \u03bc\u03b5 \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc lr. \u0391\u03c5\u03c4\u03cc \u03b4\u03b5\u03af\u03c7\u03bd\u03b5\u03b9 \u03cc\u03c4\u03b9 "
   "\u03c4\u03bf \u03b9\u03b4\u03b1\u03bd\u03b9\u03ba\u03cc lr \u03b5\u03be\u03b1\u03c1\u03c4\u03ac\u03c4\u03b1\u03b9 \u03b1\u03c0\u03cc \u03c4\u03bf backbone.")
bl(doc, "MobileNet + lr=0.005 (53.80%): \u03a4\u03bf \u03c5\u03c8\u03b7\u03bb\u03cc lr \u03c0\u03c1\u03bf\u03ba\u03b1\u03bb\u03b5\u03af \u03b1\u03c3\u03c4\u03ac\u03b8\u03b5\u03b9\u03b1 "
   "\u03c3\u03c4\u03bf \u03b5\u03bb\u03b1\u03c6\u03c1\u03cd backbone (loss \u03b1\u03c5\u03be\u03ae\u03b8\u03b7\u03ba\u03b5 \u03b1\u03c0\u03cc 0.578 \u03c3\u03b5 0.819). "
   "\u03a4\u03bf MobileNet \u03bc\u03b5 \u03bb\u03b9\u03b3\u03cc\u03c4\u03b5\u03c1\u03b5\u03c2 \u03c0\u03b1\u03c1\u03b1\u03bc\u03ad\u03c4\u03c1\u03bf\u03c5\u03c2 \u03b5\u03af\u03bd\u03b1\u03b9 \u03c0\u03b9\u03bf "
   "\u03b5\u03c5\u03b1\u03af\u03c3\u03b8\u03b7\u03c4\u03bf \u03c3\u03b5 \u03c5\u03c8\u03b7\u03bb\u03cc lr.")

im(doc, os.path.join(EX4, "report_03_best_loss_components.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 4.3: \u0391\u03bd\u03ac\u03bb\u03c5\u03c3\u03b7 loss components (classifier + box regression)")

# ── 4.4 LR and time ──────────────────────────────────────────────────────
hd(doc, "4.4 Learning Rate \u03ba\u03b1\u03b9 \u03a7\u03c1\u03cc\u03bd\u03bf\u03c2", 2)
im(doc, os.path.join(EX4, "report_04_lr_vs_accuracy.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 4.4: Learning rate vs detection accuracy")

im(doc, os.path.join(EX4, "report_05_time_vs_accuracy.png"), Inches(4.5),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 4.5: \u03a7\u03c1\u03cc\u03bd\u03bf\u03c2 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7\u03c2 vs detection accuracy")

doc.add_paragraph(
    "\u03a4\u03bf \u03b9\u03b4\u03b1\u03bd\u03b9\u03ba\u03cc lr \u03b5\u03be\u03b1\u03c1\u03c4\u03ac\u03c4\u03b1\u03b9 \u03b1\u03c0\u03cc \u03c4\u03bf backbone: "
    "\u03c4\u03bf ResNet50-FPN \u03c7\u03c1\u03b5\u03b9\u03ac\u03b6\u03b5\u03c4\u03b1\u03b9 \u03c5\u03c8\u03b7\u03bb\u03cc\u03c4\u03b5\u03c1\u03bf lr (0.005) "
    "\u03b3\u03b9\u03b1 \u03bd\u03b1 \u03c0\u03c1\u03bf\u03c3\u03b1\u03c1\u03bc\u03cc\u03c3\u03b5\u03b9 \u03c4\u03b1 \u03c0\u03bf\u03bb\u03bb\u03ac \u03c3\u03c4\u03c1\u03ce\u03bc\u03b1\u03c4\u03b1 \u03c3\u03b5 4 epochs, "
    "\u03b5\u03bd\u03ce \u03c4\u03bf MobileNet \u03bb\u03b5\u03b9\u03c4\u03bf\u03c5\u03c1\u03b3\u03b5\u03af \u03ba\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03b1 \u03bc\u03b5 \u03c7\u03b1\u03bc\u03b7\u03bb\u03cc (0.001).")

# ── 4.5 Detection results ────────────────────────────────────────────────
hd(doc, "4.5 \u0395\u03bd\u03b4\u03b5\u03b9\u03ba\u03c4\u03b9\u03ba\u03ac \u0391\u03c0\u03bf\u03c4\u03b5\u03bb\u03ad\u03c3\u03bc\u03b1\u03c4\u03b1 \u0391\u03bd\u03af\u03c7\u03bd\u03b5\u03c5\u03c3\u03b7\u03c2", 2)
im(doc, os.path.join(EX4, "report_06_best_detection_panel.png"), Inches(5.2),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 4.6: \u039a\u03b1\u03bb\u03cd\u03c4\u03b5\u03c1\u03bf detection (ResNet50, lr=0.005)")

im(doc, os.path.join(EX4, "resnet50_lr0.005", "detections.png"), Inches(5.0),
   "\u0395\u03b9\u03ba\u03cc\u03bd\u03b1 4.7: \u0394\u03b5\u03af\u03b3\u03bc\u03b1\u03c4\u03b1 detections \u2014 ResNet50, lr=0.005")

# ── 4.6 Conclusions ──────────────────────────────────────────────────────
hd(doc, "4.6 \u03a3\u03c5\u03bc\u03c0\u03b5\u03c1\u03ac\u03c3\u03bc\u03b1\u03c4\u03b1 \u0386\u03c3\u03ba\u03b7\u03c3\u03b7\u03c2 4", 2)
bl(doc, "\u03a4\u03bf Faster R-CNN \u03bc\u03b5 ResNet50-FPN \u03c0\u03b5\u03c4\u03c5\u03c7\u03b1\u03af\u03bd\u03b5\u03b9 83.18% detection accuracy "
   "\u03c3\u03b5 \u03bc\u03cc\u03bb\u03b9\u03c2 4 epochs \u2014 \u03b7 \u03b4\u03cd\u03bd\u03b1\u03bc\u03b7 \u03c4\u03bf\u03c5 pretrained COCO backbone.")
bl(doc, "\u03a4\u03bf backbone \u03ba\u03b1\u03b8\u03bf\u03c1\u03af\u03b6\u03b5\u03b9 \u03c4\u03bf \u03b9\u03b4\u03b1\u03bd\u03b9\u03ba\u03cc lr: ResNet50 \u03b8\u03ad\u03bb\u03b5\u03b9 0.005, "
   "MobileNet \u03b8\u03ad\u03bb\u03b5\u03b9 0.001. \u0397 \u03b1\u03bd\u03c4\u03b9\u03c3\u03c4\u03c1\u03bf\u03c6\u03ae τ\u03bf\u03c5\u03c2 \u03c1\u03af\u03c7\u03bd\u03b5\u03b9 \u03c3\u03b7\u03bc\u03b1\u03bd\u03c4\u03b9\u03ba\u03ac \u03c4\u03b7\u03bd \u03b1\u03c0\u03cc\u03b4\u03bf\u03c3\u03b7.")
bl(doc, "\u03a4\u03bf MobileNet \u03b5\u03af\u03bd\u03b1\u03b9 \u03ba\u03b1\u03bb\u03ae \u03b5\u03c0\u03b9\u03bb\u03bf\u03b3\u03ae \u03b3\u03b9\u03b1 \u03b3\u03c1\u03ae\u03b3\u03bf\u03c1\u03b7 \u03b5\u03ba\u03c0\u03b1\u03af\u03b4\u03b5\u03c5\u03c3\u03b7 (73.78% \u03bc\u03b5 lr=0.001), "
   "\u03b1\u03bb\u03bb\u03ac \u03c4\u03bf ResNet50-FPN \u03c5\u03c0\u03b5\u03c1\u03c4\u03b5\u03c1\u03b5\u03af \u03ba\u03b1\u03c4\u03ac 10% \u03bc\u03b5 \u03c3\u03c9\u03c3\u03c4\u03cc lr.")
bl(doc, "\u03a3\u03b5 \u03c3\u03cd\u03b3\u03ba\u03c1\u03b9\u03c3\u03b7 \u03bc\u03b5 \u03c4\u03b7\u03bd \u0386\u03c3\u03ba\u03b7\u03c3\u03b7 2 (classification 90.87%), \u03c4\u03bf detection "
   "\u03b5\u03af\u03bd\u03b1\u03b9 \u03c0\u03b9\u03bf \u03b4\u03cd\u03c3\u03ba\u03bf\u03bb\u03bf (83.18%) \u03b3\u03b9\u03b1\u03c4\u03af \u03b1\u03c0\u03b1\u03b9\u03c4\u03b5\u03af \u03ba\u03b1\u03b9 localization.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# EXERCISE 5
# ══════════════════════════════════════════════════════════════════════════
hd(doc, "Άσκηση #5: Συγκριτική CNN vs Vision Transformers")

hd(doc, "5.1 Εισαγωγή", 2)
doc.add_paragraph(
    "Συγκριτική αξιολόγηση CNN (ResNet18, 11.2M params) vs Vision Transformer "
    "(ViT-Tiny, 5.5M params) στο CIFAR-10 (10 κλάσεις). "
    "Το ResNet18 χρησιμοποιεί residual connections και συνέλιξη, "
    "ενώ το ViT χρησιμοποιεί self-attention σε image patches. "
    "4 πειράματα (cosine scheduler):")
bl(doc, "CNN ResNet18: Adam, 10 epochs, batch=128, image 32×32")
bl(doc, "ViT Tiny: AdamW, 12 epochs, batch=64, image 96×96 (upscaled)")
bl(doc, "Learning rates: 0.001, 0.0005 για κάθε αρχιτεκτονική")

hd(doc, "5.2 Συνοπτικός Πίνακας Αποτελεσμάτων", 2)
# Load ex5 results
ex5_data = []
for dn in ["cnn_resnet18_lr0.001", "cnn_resnet18_lr0.0005", "vit_tiny_lr0.0005", "vit_tiny_lr0.001"]:
    jp = os.path.join(EX5, dn, "results.json")
    if os.path.exists(jp):
        with open(jp, encoding="utf-8") as f:
            ex5_data.append(json.load(f))
cols5 = ["Πείραμα", "Αρχιτ.", "Params(M)", "LR", "Optimizer", "Test Acc(%)", "Test Loss", "Χρόνος(min)"]
rows5 = []
for e in ex5_data:
    m = e["metrics"]
    c = e["config"]
    rows5.append([e["name"], c["architecture"].upper(), f'{m["parameters"]["total_millions"]:.1f}',
                  c["learning_rate"], c["optimizer"].upper(),
                  f'{m["test_acc"]:.2f}', f'{m["test_loss"]:.4f}',
                  f'{m["total_training_time"]/60:.1f}'])
rows5.sort(key=lambda r: float(r[5]), reverse=True)
tbl(doc, cols5, rows5)
doc.add_paragraph("")

im(doc, os.path.join(EX5, "report_01_accuracy_overview.png"), Inches(4.8),
   "Εικόνα 5.1: Επισκόπηση ακρίβειας CNN vs ViT")

im(doc, os.path.join(EX5, "report_02_family_mean_accuracy.png"), Inches(4.5),
   "Εικόνα 5.2: Μέση ακρίβεια ανά οικογένεια αρχιτεκτονικής")

# ── 5.3 CNN vs ViT ──────────────────────────────────────────────────────
hd(doc, "5.3 Σύγκριση CNN vs ViT", 2)

im(doc, os.path.join(EX5, "report_03_best_cnn_vs_vit_loss.png"), Inches(4.8),
   "Εικόνα 5.3: Καμπύλες loss — Καλύτερο CNN vs Καλύτερο ViT")

doc.add_paragraph("Ανάλυση:")
bl(doc, "CNN ResNet18 (89.67%) >> ViT Tiny (62.47%): Διαφορά 27%. "
   "Οι CNN υπερτερούν σταθερά σε μικρά datasets (όπως CIFAR-10 με 50K εικόνες) "
   "γιατί έχουν ενσωματωμένο inductive bias (locality, translation equivariance).")
bl(doc, "Οι ViT χρειάζονται πολύ περισσότερα δεδομένα (π.χ. ImageNet-21K, JFT-300M) "
   "για να μάθουν τα πατρόν που οι CNN έχουν hardcoded (convolution filters). "
   "Με 50K εικόνες και 12 epochs, ο ViT δεν προλαβαίνει να συγκλίνει.")
bl(doc, "Test loss: CNN=0.31 vs ViT=1.03 — το ViT είναι πολύ λιγότερο confident "
   "στις προβλέψεις του, με υψηλότερη αβεβαιότητα σε όλες τις κλάσεις.")
bl(doc, "Χρόνος: CNN ~9.4 min vs ViT ~12.3 min (παρά τις λιγότερες παραμέτρους). "
   "Ο ViT είναι πιο αργός λόγω O(n²) self-attention και μεγαλύτερου input (96×96 vs 32×32).")

# ── 5.4 LR effect ──────────────────────────────────────────────────────
hd(doc, "5.4 Επίδραση Learning Rate", 2)
im(doc, os.path.join(EX5, "report_04_lr_vs_accuracy_by_family.png"), Inches(4.5),
   "Εικόνα 5.4: LR vs accuracy ανά οικογένεια")

doc.add_paragraph("Ανάλυση:")
bl(doc, "CNN: Ελάχιστη διαφορά μεταξύ lr=0.001 (89.67%) και lr=0.0005 (89.41%). "
   "Το ResNet18 είναι ρομπούστο σε αυτό το εύρος lr με Adam+cosine.")
bl(doc, "ViT: Επίσης μικρή διαφορά (62.47% vs 60.29%). Το lr=0.0005 είναι ελαφρώς "
   "καλύτερο, πιθανώς γιατί ο ViT είναι ευαίσθητος στο lr και χρειάζεται πιο αργή σύγκλιση.")
bl(doc, "Συμπέρασμα: Το lr δεν είναι ο κύριος παράγοντας \u2014 η αρχιτεκτονική και τα δεδομένα "
   "καθορίζουν την απόδοση περισσότερο από το fine-tuning του lr.")

# ── 5.5 Params vs accuracy ──────────────────────────────────────────────
hd(doc, "5.5 Παράμετροι vs Ακρίβεια", 2)
im(doc, os.path.join(EX5, "report_05_params_vs_accuracy.png"), Inches(4.5),
   "Εικόνα 5.5: Παράμετροι vs accuracy")

doc.add_paragraph(
    "Το ResNet18 με 11.2M params πετυχαίνει 89.67%, ενώ ο ViT με 5.5M params "
    "μόλις 62.47%. Περισσότερες παράμετροι δεν είναι κατ' ανάγκη καλύτερες, "
    "αλλά στην περίπτωση αυτή το inductive bias των CNN αντισταθμίζει "
    "την έλλειψη δεδομένων. Σε μεγάλα datasets, ο ViT θα υπερτερούσε.")

# ── 5.6 Sample predictions ──────────────────────────────────────────────
hd(doc, "5.6 Ενδεικτικά Αποτελέσματα Ταξινόμησης", 2)
im(doc, os.path.join(EX5, "report_06_best_predictions_panel.png"), Inches(5.2),
   "Εικόνα 5.6: Καλύτερες προβλέψεις CNN vs ViT")

im(doc, os.path.join(EX5, "all_training_curves.png"), Inches(5.0),
   "Εικόνα 5.7: Καμπύλες εκπαίδευσης όλων των πειραμάτων")

# ── 5.7 Conclusions ─────────────────────────────────────────────────────
hd(doc, "5.7 Συμπεράσματα Άσκησης 5", 2)
bl(doc, "Οι CNN (π.χ. ResNet18) υπερτερούν σημαντικά των ViT (89.67% vs 62.47%) σε "
   "μικρά datasets λόγω inductive bias (locality, weight sharing).")
bl(doc, "Οι ViT χρειάζονται περισσότερα δεδομένα και epochs για να συγκλίνουν. "
   "Με pretrained ViT (π.χ. από ImageNet-21K), η διαφορά αναμένεται να εξαλειφθεί.")
bl(doc, "Το lr έχει μικρή επίδραση στο εύρος 0.0005-0.001. "
   "Η αρχιτεκτονική/δεδομένα κυριαρχούν σε αυτό το scenario.")
bl(doc, "Σε σύγκριση με Άσκηση 1 (CNN στο CIFAR-100: 56.43%), "
   "το ResNet18 στο CIFAR-10 φτάνει 89.67% \u2014 10 κλάσεις vs 100 κάνουν τεράστια διαφορά.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# GENERAL CONCLUSIONS
# ══════════════════════════════════════════════════════════════════════════
hd(doc, "Γενικά Συμπεράσματα")
bl(doc, "Άσκηση 1: Vanilla CNN με SGD+cosine πετυχαίνει 56.43% στο CIFAR-100. Ο ρυθμός μάθησης είναι η πιο κρίσιμη υπερ-παράμετρος.")
bl(doc, "Άσκηση 2: Transfer learning φτάνει 90.87% με ResNet50 frozen. Τα pretrained features είναι εξαιρετικά ισχυρά.")
bl(doc, "Άσκηση 3: U-Net πετυχαίνει mIoU=13.91% με κανονικό BG weight. Η μείωση του BG weight απαιτεί περισσότερες epochs.")
bl(doc, "Άσκηση 4: Faster R-CNN με ResNet50-FPN πετυχαίνει 83.18% detection. "
   "Το backbone καθορίζει το ιδανικό lr.")
bl(doc, "Άσκηση 5: CNN (89.67%) >> ViT (62.47%) στο CIFAR-10. "
   "Οι transformers χρειάζονται πολύ περισσότερα δεδομένα.")
bl(doc, "Γενικά: Η επιλογή αρχιτεκτονικής, optimizer και lr εξαρτάται από το task και τα δεδομένα. "
   "CNN για μικρά datasets, ViT για μεγάλα. Transfer learning είναι πάντα αποτελεσματικό.")

doc.save(OUTPUT)
print(f"\n[OK] Saved: {OUTPUT}")
print(f"  Size: {os.path.getsize(OUTPUT)/(1024*1024):.1f} MB")
